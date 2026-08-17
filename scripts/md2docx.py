#!/usr/bin/env python3
"""md2docx — 将结构化 markdown 重建为 .docx (格式转换层: 输出阶段).

与 docx2md.py 互逆: md 只承载内容与结构 (段落/标题层级/脚注),
排版由本脚本按规则表统一套用 — 与 format_paper.format_document
的排版规则完全一致 (题目黑体四号居中 / 一级宋体小四加粗居中 /
二级楷体小四加粗缩进2 / 三级宋体五号加粗缩进2 / 摘要关键词楷体小四 /
正文宋体五号缩进2 / 脚注小五号宋体).

Markdown 约定 (与 docx2md.py 一致):
  # 题目        → 论文题目
  (无标记段落)   → 作者行 / 副标题 / 摘要 / 关键词
  ## 一、xxx    → 一级标题
  ### （一）xxx → 二级标题
  #### 1. xxx   → 三级标题
  [^n]          → 脚注引用 (段落文本中的标记, 重建时转为脚注引用上标)
  [^n]: 文本    → 脚注定义

用法:
  python md2docx.py 论文.md -o 论文.docx
  python md2docx.py 论文.md            # 输出到同名 .docx
"""

import argparse
import io
import re
import sys
import zipfile
from pathlib import Path
from typing import Dict, List, Tuple

from docx import Document
from docx.shared import Pt
from lxml import etree

from format_paper import (FONT_EN, _ALIGN_MAP,
                          apply_run_font, detect_heading_level,
                          is_abstract_or_keywords,
                          remove_first_line_indent, set_alignment,
                          set_first_line_indent, set_line_spacing)

# OOXML 命名空间
WML = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
REL = 'http://schemas.openxmlformats.org/package/2006/relationships'
FOOTNOTE_REL = ('http://schemas.openxmlformats.org/officeDocument/'
                '2006/relationships/footnotes')
CT = 'http://schemas.openxmlformats.org/package/2006/content-types'

_HEADING_RE = re.compile(r'^(#{1,5})\s+(.*)$')
_FN_DEF_RE = re.compile(r'^\[\^(\d+)\]:\s*(.*)$')
_FN_REF_RE = re.compile(r'\[\^(\d+)\]')


# ── Markdown 解析 ─────────────────────────────────────────────────────


def parse_markdown(md_text: str) -> Tuple[List[dict], Dict[int, str]]:
    """解析 markdown 为有序段落块与脚注定义.

    Returns:
        (blocks, footnotes):
        blocks: [{'kind': 'title'|'author'|'abstract'|'keywords'|
                          'h1'|'h2'|'h3'|'body',
                  'text': str}]
        footnotes: {id: text}
    """
    blocks: List[dict] = []
    footnotes: Dict[int, str] = {}
    current = None       # 正在合并的普通段落
    fn_state = None      # 正在收集续行的脚注 id

    def flush():
        nonlocal current
        if current is not None and current['text']:
            blocks.append(current)
        current = None

    for raw_line in md_text.splitlines():
        line = raw_line.strip()
        if not line:
            flush()
            fn_state = None
            continue

        # 新脚注定义行: 先于续行检查 (定义间可能无空行分隔)
        m = _FN_DEF_RE.match(line)
        if m:
            flush()
            fn_state = int(m.group(1))
            footnotes[fn_state] = m.group(2).strip()
            continue

        # 脚注定义续行
        if fn_state is not None:
            footnotes[fn_state] = footnotes[fn_state] + ' ' + line
            continue

        m = _HEADING_RE.match(line)
        if m:
            flush()
            kind = {1: 'title', 2: 'h1', 3: 'h2', 4: 'h3', 5: 'h4'}[len(m.group(1))]
            blocks.append({'kind': kind, 'text': m.group(2).strip()})
            continue

        # 普通段落 (连续行合并)
        if current is None:
            current = {'kind': 'body', 'text': ''}
        current['text'] += line
    flush()

    # 作者行/摘要/关键词分类: 题目之后、第一个标题之前的普通段
    seen_title = False
    seen_heading = False
    for b in blocks:
        if b['kind'] == 'title':
            seen_title = True
            continue
        if b['kind'] in ('h1', 'h2', 'h3', 'h4'):
            seen_heading = True
            continue
        if b['kind'] == 'body':
            if seen_title and not seen_heading:
                ak = is_abstract_or_keywords(b['text'])
                if ak == 'abstract':
                    b['kind'] = 'abstract'
                elif ak == 'keywords':
                    b['kind'] = 'keywords'
                else:
                    b['kind'] = 'author'  # 作者行/副标题

    return blocks, footnotes


def _split_fn_refs(text: str) -> Tuple[str, List[int]]:
    """从段落文本中提取 [^n] 引用标记, 返回 (清理后文本, id 列表)."""
    ids = [int(m.group(1)) for m in _FN_REF_RE.finditer(text)]
    cleaned = _FN_REF_RE.sub('', text).strip()
    return cleaned, ids


# ── 排版应用 (规则表驱动, 与 format_paper.format_document 一致) ────────


def _apply_format(doc, text: str, kind: str, body_indent: int,
                  fn_ids: List[int], rules: dict):
    """按规则表应用排版, 返回段落对象.

    kind 即规则元素键: title/abstract/keywords/author/h1/h2/h3/h4/body.
    """
    rule = rules[kind]
    para = doc.add_paragraph()

    set_alignment(para, _ALIGN_MAP[rule['align']])
    set_line_spacing(para, rule['line_spacing'])
    indent = rule['indent']
    if indent is None:
        indent = body_indent
    if indent and indent > 0:
        set_first_line_indent(para, chars=indent)
    else:
        remove_first_line_indent(para)
    run = para.add_run(text)
    apply_run_font(run, rule['font'], FONT_EN,
                   Pt(rule['size']), bold=rule['bold'])

    # 脚注引用: 段末追加上标引用 run
    for fn_id in fn_ids:
        r_elem = etree.SubElement(para._p, f'{{{WML}}}r')
        rPr = etree.SubElement(r_elem, f'{{{WML}}}rPr')
        va = etree.SubElement(rPr, f'{{{WML}}}vertAlign')
        va.set(f'{{{WML}}}val', 'superscript')
        fn_ref = etree.SubElement(r_elem, f'{{{WML}}}footnoteReference')
        fn_ref.set(f'{{{WML}}}id', str(fn_id))

    return para


# ── 脚注注入 (zipfile 模式, 仿 tests/conftest._inject_footnotes) ──────


def _build_footnotes_xml(footnotes: Dict[int, str], rule: dict = None) -> bytes:
    """构造 footnotes.xml: 分隔符 + 真实脚注 (上标 [n] + 文本).

    Args:
        footnotes: {id: 文本}.
        rule: footnote 元素规则 (load_rules 输出); None = 默认 (宋体 9pt 单倍).
    格式模式与 format_paper._format_footnote_paragraph 一致 (rFonts/sz + 行距).
    """
    if rule is None:
        rule = {'font': '宋体', 'size': 9, 'line_spacing': 1.0}
    fns_root = etree.Element(f'{{{WML}}}footnotes', nsmap={'w': WML})

    # 分隔符脚注 (id 0/-1), 与 Word 生成的文档一致
    # 必须带 w:type: 无 type 时 Word 按普通脚注解析, id=-1 非法 → 拒开文档
    for sep_id, sep_tag in [('-1', 'separator'), ('0', 'continuationSeparator')]:
        sep = etree.SubElement(fns_root, f'{{{WML}}}footnote')
        sep.set(f'{{{WML}}}type', sep_tag)
        sep.set(f'{{{WML}}}id', sep_id)
        sp = etree.SubElement(sep, f'{{{WML}}}p')
        sr = etree.SubElement(sp, f'{{{WML}}}r')
        etree.SubElement(sr, f'{{{WML}}}{sep_tag}')

    size_hp = str(int(rule['size'] * 2))
    for fn_id in sorted(footnotes):
        fn = etree.SubElement(fns_root, f'{{{WML}}}footnote')
        fn.set(f'{{{WML}}}id', str(fn_id))
        p = etree.SubElement(fn, f'{{{WML}}}p')

        # 段落行距 (与 _format_footnote_paragraph 相同逻辑)
        ls = rule['line_spacing']
        if isinstance(ls, dict) and ls['mode'] == 'exact':
            line, line_rule = str(int(ls['value'] * 20)), 'exact'
        else:
            line, line_rule = str(int(float(ls) * 240)), 'auto'
        pPr = etree.SubElement(p, f'{{{WML}}}pPr')
        spacing = etree.SubElement(pPr, f'{{{WML}}}spacing')
        spacing.set(f'{{{WML}}}line', line)
        spacing.set(f'{{{WML}}}lineRule', line_rule)
        spacing.set(f'{{{WML}}}before', '0')
        spacing.set(f'{{{WML}}}after', '0')

        def _mk_run(text=None):
            r = etree.SubElement(p, f'{{{WML}}}r')
            rPr = etree.SubElement(r, f'{{{WML}}}rPr')
            rFonts = etree.SubElement(rPr, f'{{{WML}}}rFonts')
            rFonts.set(f'{{{WML}}}eastAsia', rule['font'])
            rFonts.set(f'{{{WML}}}ascii', rule['font'])
            rFonts.set(f'{{{WML}}}hAnsi', rule['font'])
            sz = etree.SubElement(rPr, f'{{{WML}}}sz')
            sz.set(f'{{{WML}}}val', size_hp)
            szCs = etree.SubElement(rPr, f'{{{WML}}}szCs')
            szCs.set(f'{{{WML}}}val', size_hp)
            if text is not None:
                t = etree.SubElement(r, f'{{{WML}}}t')
                t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
                t.text = text
            return r

        # 左括号
        _mk_run('[')
        # 自动编号
        r1 = _mk_run()
        etree.SubElement(r1, f'{{{WML}}}footnoteRef')
        # 右括号 + 文本
        _mk_run(']' + footnotes[fn_id])

    return etree.tostring(fns_root, xml_declaration=True,
                          encoding='UTF-8', standalone=True)


def _inject_footnotes(docx_bytes: bytes, footnotes: Dict[int, str],
                      rule: dict = None) -> bytes:
    """向 .docx ZIP 注入 footnotes.xml + rels + content-types 覆盖.

    与 tests/conftest._inject_footnotes 同模式, 但避免
    [Content_Types].xml 被重复写入 (循环中跳过, 末尾统一写).
    rule: footnote 元素规则; None = 默认 (宋体 9pt 单倍).
    """
    with zipfile.ZipFile(io.BytesIO(docx_bytes), 'r') as zin:
        out_buf = io.BytesIO()
        with zipfile.ZipFile(out_buf, 'w', zipfile.ZIP_DEFLATED) as zout:
            for item in zin.infolist():
                if item.filename in ('[Content_Types].xml', 'word/footnotes.xml'):
                    continue  # 末尾统一写 (footnotes.xml 同样跳过, 避免 zip 重复条目)
                data = zin.read(item.filename)

                if item.filename == 'word/_rels/document.xml.rels':
                    rels_xml = etree.fromstring(data)
                    # 查重: 已存在同类型关系时复用, 避免重复 rel (Word 判损坏)
                    existing = [r for r in rels_xml
                                if r.get('Type') == FOOTNOTE_REL]
                    if not existing:
                        fn_rel = etree.SubElement(rels_xml, f'{{{REL}}}Relationship')
                        fn_rel.set('Id', 'rIdFootnotes')
                        fn_rel.set('Type', FOOTNOTE_REL)
                        fn_rel.set('Target', 'footnotes.xml')
                    data = etree.tostring(rels_xml, xml_declaration=True,
                                          encoding='UTF-8', standalone=True)

                zout.writestr(item, data)

            zout.writestr('word/footnotes.xml', _build_footnotes_xml(footnotes, rule))

            # Content_Types: 统一在末尾写入 (避免重复条目)
            ct_xml = etree.fromstring(zin.read('[Content_Types].xml'))
            override = etree.SubElement(ct_xml, f'{{{CT}}}Override')
            override.set('PartName', '/word/footnotes.xml')
            override.set('ContentType', ('application/vnd.openxmlformats-'
                                         'officedocument.wordprocessingml.'
                                         'footnotes+xml'))
            zout.writestr('[Content_Types].xml',
                          etree.tostring(ct_xml, xml_declaration=True,
                                         encoding='UTF-8', standalone=True))

    return out_buf.getvalue()


# ── 主流程 ────────────────────────────────────────────────────────────


def markdown_to_docx(md_text: str, output_path: str,
                     body_indent: int = 2, rules: dict = None) -> dict:
    """将 markdown 重建为 .docx.

    Args:
        md_text: 结构化 markdown 文本.
        output_path: 输出 .docx 路径.
        body_indent: 正文字符缩进数 (默认2, 0=不缩进).
        rules: 自定义格式规则表 (rules.py load_rules 输出);
               None = 使用默认规则 (skill 预设格式).

    Returns:
        统计信息 dict: {title, headings_l1..l4, body, footnotes, author}
    """
    if rules is None:
        from rules import DEFAULT_RULES
        rules = DEFAULT_RULES

    blocks, footnotes = parse_markdown(md_text)
    if not any(b['kind'] == 'title' for b in blocks):
        raise ValueError('未检测到题目: markdown 缺少 "# " 首行标题')

    stats = {'title': 0, 'author': 0, 'abstract': 0, 'keywords': 0,
             'headings_l1': 0, 'headings_l2': 0, 'headings_l3': 0,
             'headings_l4': 0, 'body': 0, 'footnotes': len(footnotes)}

    doc = Document()
    referenced: List[int] = []
    for b in blocks:
        text, fn_ids = _split_fn_refs(b['text'])
        referenced.extend(fn_ids)
        _apply_format(doc, text, b['kind'], body_indent, fn_ids, rules)
        stats['title' if b['kind'] == 'title'
              else 'author' if b['kind'] == 'author'
              else 'abstract' if b['kind'] == 'abstract'
              else 'keywords' if b['kind'] == 'keywords'
              else f'headings_l{b["kind"][1]}' if b['kind'].startswith('h')
              else 'body'] += 1

    # 引用了但未定义的脚注 → 补齐空定义 (避免 Word 打开报错)
    for fn_id in referenced:
        if fn_id not in footnotes:
            footnotes[fn_id] = ''
    stats['footnotes'] = len(footnotes)

    # —— 标点规范化 (正文/摘要等, 跳过题目与各级标题) ——
    try:
        from citation_formatter import normalize_chinese_punctuation
        punct_fixes = 0
        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            if not text:
                continue
            if i == 0 or detect_heading_level(text) > 0:
                continue  # 题目/各级标题: 不修改标点
            for run in para.runs:
                if run.text:
                    new_text = normalize_chinese_punctuation(run.text)
                    if new_text != run.text:
                        punct_fixes += 1
                        run.text = new_text
        if punct_fixes > 0:
            stats['punct_fixes'] = punct_fixes
    except ImportError:
        pass

    buf = io.BytesIO()
    doc.save(buf)

    if footnotes:
        data = _inject_footnotes(buf.getvalue(), footnotes, rules['footnote'])
    else:
        data = buf.getvalue()

    Path(output_path).write_bytes(data)
    return stats


def build_parser() -> argparse.ArgumentParser:
    ap = argparse.ArgumentParser(
        description='markdown → docx 重建 (输出阶段: 排版按规则表统一套用)')
    ap.add_argument('input', help='输入 .md 文件')
    ap.add_argument('-o', '--output', default=None,
                    help='输出 .docx (默认与输入同名)')
    ap.add_argument('--body-indent', type=int, default=2,
                    help='正文首行缩进字符数 (默认 2, 0=不缩进)')
    ap.add_argument('--rules', default=None,
                    help='自定义格式规则 JSON 配置文件')
    return ap


def main():
    args = build_parser().parse_args()
    inp = Path(args.input)
    if not inp.exists():
        print(f"错误: 输入文件不存在: {inp}", file=sys.stderr)
        sys.exit(1)
    if inp.suffix.lower() != '.md':
        print(f"错误: 不支持的文件格式: {inp.suffix} (仅支持 .md)", file=sys.stderr)
        sys.exit(1)

    output = args.output or str(inp.with_suffix('.docx'))
    md_text = inp.read_text(encoding='utf-8')
    rules = None
    if args.rules:
        from rules import load_rules
        try:
            rules = load_rules(args.rules)
            print(f"已加载自定义格式规则: {args.rules}")
        except (ValueError, FileNotFoundError) as e:
            print(f"错误: {e}", file=sys.stderr)
            sys.exit(1)
    stats = markdown_to_docx(md_text, output, body_indent=args.body_indent,
                             rules=rules)
    print(f"已重建: {output}")
    parts = []
    for k, label in [('title', '题目'), ('author', '作者行'),
                     ('abstract', '摘要'), ('keywords', '关键词'),
                     ('headings_l1', '一级标题'), ('headings_l2', '二级标题'),
                     ('headings_l3', '三级标题'), ('body', '正文'),
                     ('footnotes', '脚注')]:
        if stats.get(k):
            parts.append(f"{label}: {stats[k]}")
    if stats.get('punct_fixes'):
        parts.append(f"标点修复: {stats['punct_fixes']}处")
    print(f"✓ 完成 — {' | '.join(parts)}")


if __name__ == '__main__':
    main()
