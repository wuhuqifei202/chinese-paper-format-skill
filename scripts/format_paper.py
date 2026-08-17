#!/usr/bin/env python3
"""
Chinese Academic Paper Formatter — 中文学术论文格式规范化工具

按照中文学术论文排版规范，自动检测并设置 .docx 文件的格式：
  论文题目: 黑体, 四号(14pt), 不加粗, 居中
  一级标题 (一、): 宋体, 小四(12pt), 加粗, 单倍行距, 居中
  二级标题 (（一）): 楷体, 小四(12pt), 加粗, 单倍行距, 缩进2字符
  三级标题 (1. ): 宋体, 五号(10.5pt), 加粗, 单倍行距, 缩进2字符
  正文: 宋体(中文+数字) + Times New Roman(西文字母), 五号(10.5pt), 单倍行距
  脚注: 宋体, 小五号(9pt), 单倍行距, 每页重新编号

Usage:
    python format_paper.py input.docx                  # 覆盖原文件
    python format_paper.py input.docx -o output.docx   # 输出到新文件
    python format_paper.py input.docx --check          # 仅检查不修改
    python format_paper.py folder/ --batch             # 批量处理
"""

import argparse
import os
import re
import sys
import shutil
from datetime import datetime
from pathlib import Path
from typing import List, Optional, Tuple

# 确保脚本目录在 sys.path 中 — 裸 import citation_rules / citation_formatter 依赖于此
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from docx.shared import Pt
from lxml import etree

# 引注格式化模块 (可选导入)
try:
    from citation_formatter import (extract_footnotes, check_footnote,
                                     format_all_footnotes)
    _HAS_CITATION = True
except ImportError:
    _HAS_CITATION = False

# 注释体例模块 (可选导入)
try:
    from citation_rules import load_citation_rules as _load_citation_rules
    _HAS_CITATION_RULES = True
except ImportError:
    _HAS_CITATION_RULES = False

# ---------------------------------------------------------------------------
# 字号常量 (单位: Pt)
# ---------------------------------------------------------------------------
SIZE_SIHAO = Pt(14)       # 四号    — 论文题目
SIZE_XIAOSI = Pt(12)      # 小四    — 一级标题 / 二级标题
SIZE_WUHAO = Pt(10.5)     # 五号    — 三级标题 / 正文
SIZE_XIAOWU = Pt(9)       # 小五号  — 脚注

# ---------------------------------------------------------------------------
# 字体常量
# ---------------------------------------------------------------------------
FONT_HEI = '黑体'                # 论文题目
FONT_SONG = '宋体'               # 一级/三级标题, 正文, 脚注
FONT_KAI = '楷体'                # 二级标题
FONT_EN = '宋体'                # 西文/数字字体

# ---------------------------------------------------------------------------
# 中文数字结构化验证
# ---------------------------------------------------------------------------
# 不再用贪心字符类 [一二三...]+，而是显式匹配合法的中文章节数字结构。
# 合法模式 (| 分隔):
#   一~九  十~十九  二十~九十九  一百~九百  一百零一~九百零九
#   一百一十~九百九十  一百十一~九百九十九
# 非法示例: 十三五 (十后跟两个数字, 不符合任何结构)
_VALID_CN_NUMBER = re.compile(
    r'^('
    r'[一二三四五六七八九]'                                # 1-9
    r'|十[一二三四五六七八九]?'                             # 10-19
    r'|[一二三四五六七八九]十[一二三四五六七八九]?'         # 20-99
    r'|[一二三四五六七八九]百'                              # 100-900
    r'|[一二三四五六七八九]百零[一二三四五六七八九]'        # 101-109
    r'|[一二三四五六七八九]百[一二三四五六七八九]十'        # 110-990, no ones
    r'|[一二三四五六七八九]百[一二三四五六七八九]十[一二三四五六七八九]'  # 111-999
    r')$'
)

# 一级标题: 中文数字 + 分隔符 (用结构化验证替换字符类)
_RE_LEVEL1 = re.compile(r'^([一二三四五六七八九十百千零]+)([、．.])')

# 二级标题: （中文数字）
_RE_LEVEL2 = re.compile(r'^（([一二三四五六七八九十百千零]+)）')

# 三级标题: 1.  2.  3.  …  或 1、 2、 … 或 1． …
_RE_LEVEL3 = re.compile(r'^\d{1,3}[\.、．]\s')

# 疑似三级标题但后面没有空格的情况 (1.xxx → 可能是序号)
_RE_LEVEL3_LOOSE = re.compile(r'^\d{1,3}[\.、．]\S')

# 四级标题: （1）（2）… 阿拉伯数字圆括号
_RE_LEVEL4 = re.compile(r'^（(\d{1,2})）')

# 摘要/关键词检测
_RE_ABSTRACT = re.compile(r'^[【\[]\s*摘\s*要\s*[】\]]|^摘要[：:)]|^摘要\b')
_RE_KEYWORDS = re.compile(r'^[【\[]\s*关\s*键\s*词\s*[】\]]|^关键词[：:)]|^关键词\b')

# 无编号一级标题: 引言、绪论、前言、结论、结语、导论、序言、余论
_UNNUMBERED_HEADINGS = {'引言', '绪论', '前言', '结论', '结语', '导论', '序言', '余论'}


def is_abstract_or_keywords(text: str) -> Optional[str]:
    """检测段落是否为摘要或关键词.

    Returns:
        'abstract' — 摘要段落
        'keywords' — 关键词段落
        None — 都不是
    """
    if _RE_ABSTRACT.match(text):
        return 'abstract'
    if _RE_KEYWORDS.match(text):
        return 'keywords'
    return None

# ---------------------------------------------------------------------------
# 最大标题长度 (字符) — 超过此长度的行不太可能是标题
# ---------------------------------------------------------------------------
MAX_HEADING_LEN = {
    1: 40,   # 一级标题最大字符数
    2: 50,   # 二级标题最大字符数
    3: 60,   # 三级标题最大字符数
    4: 20,   # 四级标题最大字符数 (更严格, 避免与正文列表混淆)
}


def detect_heading_level(text: str) -> int:
    """检测段落文本的标题层级.

    Args:
        text: 段落文本 (已去除首尾空白).

    Returns:
        0 — 正文
        1 — 一级标题 (一、)
        2 — 二级标题 (（一）)
        3 — 三级标题 (1. )
        4 — 四级标题 (（1）)
    """
    if not text:
        return 0

    # 四级标题: （1）（2）… 阿拉伯数字圆括号 — 仅很短时判定, 避免与正文列表混淆
    m4 = _RE_LEVEL4.match(text)
    if m4 and len(text) <= MAX_HEADING_LEN[4]:
        return 4

    # 三级标题最容易误判 (普通数字开头), 用更严格的规则
    m3 = _RE_LEVEL3.match(text)
    m3_loose = _RE_LEVEL3_LOOSE.match(text)
    if m3:
        # "1. xxx" with space after dot — 高度可信
        if len(text) <= MAX_HEADING_LEN[3]:
            return 3
    elif m3_loose:
        # "1.xxx" without space — 可能是列表项, 仅在很短时判定为标题
        if len(text) <= 15:
            return 3
        return 0

    # 二级标题: （一）...
    m2 = _RE_LEVEL2.match(text)
    if m2 and len(text) <= MAX_HEADING_LEN[2]:
        cn_num = m2.group(1)  # 提取括号内的中文数字
        if _VALID_CN_NUMBER.match(cn_num):
            return 2

    # 一级标题: 一、...
    # 增加保护: "第一，" "其一、" 等不是标题
    # 增加中文数字结构化验证: "十三五、" 不是合法章节数字
    m1 = _RE_LEVEL1.match(text)
    if m1 and len(text) <= MAX_HEADING_LEN[1]:
        prefix = text[:m1.start()] if m1.start() > 0 else ''
        if prefix:
            return 0  # "第一，" "其一、" 等
        cn_num = m1.group(1)  # 提取分隔符前的中文数字
        if _VALID_CN_NUMBER.match(cn_num):
            return 1

    # 无编号一级标题: 引言、绪论、前言、结论 等
    if text in _UNNUMBERED_HEADINGS:
        return 1

    return 0


def apply_run_font(run, cn_font: str, en_font: str, size,
                   bold: bool = False, italic: bool = False):
    """对单个 Run 设置字体属性.

    python-docx 高层 API 设置中文字体有限; 此处直接操作 XML 确保
    eastAsia (中文字体) 和 ascii/hAnsi (西文字体) 同时生效.
    """
    run.font.size = size
    run.bold = bold
    run.italic = italic

    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} />')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), cn_font)
    rFonts.set(qn('w:ascii'), en_font)
    rFonts.set(qn('w:hAnsi'), en_font)
    rFonts.set(qn('w:cs'), en_font)


def set_line_spacing(para, line_spacing=1.0):
    """设置段落行距, 并清除段前/段后间距.

    Args:
        line_spacing: float = 倍数行距 (1.0 = 单倍);
                      dict {'mode': 'exact', 'value': 20} = 固定 20 磅行距.
    """
    pPr = para._p.get_or_add_pPr()
    spacing = pPr.find(qn('w:spacing'))
    if spacing is None:
        spacing = parse_xml(f'<w:spacing {nsdecls("w")} />')
        _insert_pPr_ordered(pPr, spacing)

    if isinstance(line_spacing, dict) and line_spacing.get('mode') == 'exact':
        # 固定磅值行距: 1pt = 20 twips, lineRule=exact
        spacing.set(qn('w:line'), str(int(line_spacing['value'] * 20)))
        spacing.set(qn('w:lineRule'), 'exact')
    else:
        # 倍数行距: 单倍 = 240 twips (20 twips * 12pt * 1.0)
        spacing.set(qn('w:line'), str(int(float(line_spacing) * 240)))
        spacing.set(qn('w:lineRule'), 'auto')
    # OOXML CT_Spacing: before 与 beforeLines 互斥 (EG_SpaceBefore choice),
    # after 与 afterLines 同理 — 同时存在会被 Word 判为文档损坏.
    spacing.attrib.pop(qn('w:beforeLines'), None)
    spacing.attrib.pop(qn('w:afterLines'), None)
    spacing.set(qn('w:before'), '0')
    spacing.set(qn('w:after'), '0')
    # 属性追加顺序违反 CT_Spacing schema → Word 判损坏, 需重排
    _reorder_attrs(spacing, _SPACING_ATTR_ORDER)


# OOXML schema 规定的属性顺序 (CT_Ind / CT_Spacing).
# lxml 的 Element.set() 将新属性追加到属性尾部, 与文档原有属性合并后
# 顺序可能违反 schema (如 firstLineChars 出现在 right 之后), 新版 Word
# 会判定文档损坏并忽略这些属性 → 正文首行缩进"看起来没生效".
_IND_ATTR_ORDER = ('start', 'startChars', 'end', 'endChars', 'hanging',
                   'hangingChars', 'firstLine', 'firstLineChars',
                   'left', 'leftChars', 'right', 'rightChars')
_SPACING_ATTR_ORDER = ('before', 'beforeLines', 'beforeAutospacing',
                       'after', 'afterLines', 'afterAutospacing',
                       'line', 'lineRule')


def _reorder_attrs(el, order: tuple):
    """按 OOXML schema 顺序重排元素属性.

    未列入 order 的属性保持原相对顺序追加在尾部.
    """
    attrs = dict(el.attrib)
    el.attrib.clear()
    for key in order:
        qkey = qn('w:' + key)
        if qkey in attrs:
            el.attrib[qkey] = attrs.pop(qkey)
    for key, val in attrs.items():
        el.attrib[key] = val


# CT_PPr 全序 (ECMA-376 CT_PPr), 用于 pPr 子元素插入位置判定
_CT_PPR_ORDER = ('pStyle', 'keepNext', 'keepLines', 'pageBreakBefore',
                 'framePr', 'widowControl', 'numPr', 'suppressLineNumbers',
                 'pBdr', 'shd', 'tabs', 'suppressAutoHyphens', 'kinsoku',
                 'wordWrap', 'overflowPunct', 'topLinePunct', 'autoSpaceDE',
                 'autoSpaceDN', 'bidi', 'adjustRightInd', 'snapToGrid',
                 'spacing', 'ind', 'contextualSpacing', 'mirrorIndents',
                 'suppressOverlap', 'jc', 'textDirection', 'textAlignment',
                 'textboxTightWrap', 'outlineLvl', 'divId', 'cnfStyle',
                 'rPr', 'sectPr', 'pPrChange')


def _insert_pPr_ordered(pPr, el):
    """按 CT_PPr 全序把 el 插入 pPr 正确位置.

    不能直接 pPr.append(): 若 pPr 已有 jc 等顺序靠后的元素, 新元素排到
    其后 → pPr 元素顺序违规 → Word 判文档损坏. 也不能用 python-docx 的
    get_or_add_ind()/get_or_add_spacing(): 其 successors 列表顺序与 XML
    实际顺序无关, 会把元素插到 rPr 之后.
    """
    _POS = {name: i for i, name in enumerate(_CT_PPR_ORDER)}
    self_pos = _POS[etree.QName(el).localname]
    for child in pPr:
        if _POS.get(etree.QName(child).localname, -1) > self_pos:
            child.addprevious(el)
            return
    pPr.append(el)


def set_first_line_indent(para, chars: int = 2):
    """设置首行缩进 (字符数)."""
    pPr = para._p.get_or_add_pPr()
    ind = pPr.find(qn('w:ind'))
    if ind is None:
        ind = parse_xml(f'<w:ind {nsdecls("w")} />')
        _insert_pPr_ordered(pPr, ind)

    ind.set(qn('w:firstLineChars'), str(chars * 100))
    ind.set(qn('w:leftChars'), '0')
    # firstLineChars 与 hanging/hangingChars 互斥 (EG_HangingIndent choice),
    # 同时存在会被 Word 判为文档损坏 — 必须清除 (含原文继承的 hangingChars)
    ind.attrib.pop(qn('w:hangingChars'), None)
    # 清除绝对数值, 避免冲突
    for attr in ('w:firstLine', 'w:left', 'w:hanging'):
        val = ind.get(qn(attr))
        if val is not None:
            del ind.attrib[qn(attr)]
    # 属性追加顺序违反 CT_Ind schema → Word 判损坏, 需重排
    _reorder_attrs(ind, _IND_ATTR_ORDER)


def remove_first_line_indent(para):
    """移除首行缩进."""
    pPr = para._p.get_or_add_pPr()
    ind = pPr.find(qn('w:ind'))
    if ind is not None:
        for attr in ('w:firstLineChars', 'w:firstLine',
                     'w:leftChars', 'w:left',
                     'w:hangingChars', 'w:hanging'):
            val = ind.get(qn(attr))
            if val is not None:
                del ind.attrib[qn(attr)]


def set_alignment(para, alignment):
    """设置段落对齐方式."""
    para.alignment = alignment


def clear_paragraph_runs(para):
    """若段落无 run, 插入一个空 run (确保格式可写入)."""
    if not para.runs:
        # 使用 add_run 可能导致 XML 顺序问题; 直接用 lxml 插入
        r = parse_xml(f'<w:r {nsdecls("w")}><w:t xml:space="preserve"></w:t></w:r>')
        para._p.append(r)


def _drop_unused_numbering(doc):
    """移除无引用的 numbering 部件.

    源 .doc 经 Word 转换生成的 docx 常带空 numbering.xml (无 abstractNum/num
    定义), 且 document.xml/styles.xml 中无 numPr/numId 引用. 空部件 + 关系
    会使 Word 判定"文档已损坏"。有编号引用时保留 (python-docx 不管理
    numbering 部件, 此时不处理).
    """
    try:
        refs = [doc.element.xml]
        if doc.styles is not None:
            refs.append(doc.styles.element.xml)
        if any('numPr' in t or 'numId' in t for t in refs):
            return
        for rel in list(doc.part.rels.values()):
            if rel.reltype.endswith('/numbering'):
                doc.part.drop_rel(rel.rId)
    except Exception:
        pass


def _dedupe_footnotes_rel(doc):
    """去重指向同一 footnotes part 的重复关系.

    不同阶段 (Word .doc 转换 / 引注段转脚注 / md2docx 注入) 可能各自添加一条
    footnotes 关系 (如 rId5 与 rIdFootnotes 同指 footnotes.xml). 同一 part
    存在多条指向它的关系会被 Word 判为"文档已损坏" — 仅保留一条, 并同步
    改写 document.xml 中引用被删 Id 的 r:id 属性.
    """
    FOOTNOTES_TYPE = ('http://schemas.openxmlformats.org/officeDocument/'
                      '2006/relationships/footnotes')
    fn_rels = [rel for rel in doc.part.rels.values()
               if rel.reltype == FOOTNOTES_TYPE]
    if len(fn_rels) <= 1:
        return
    keep = fn_rels[0]
    drop_ids = {rel.rId for rel in fn_rels[1:]}

    # 改写 document.xml 中对被删 Id 的引用 (r:id 位于 officeDocument
    # relationships 命名空间, 与 .rels 文件的 package 命名空间不同)
    from docx.oxml.ns import qn as _qn
    _OFF_RID = ('{http://schemas.openxmlformats.org/officeDocument/'
                '2006/relationships}id')
    for el in doc.element.iter():
        rid = el.get(_OFF_RID)
        if rid in drop_ids:
            el.set(_OFF_RID, keep.rId)

    for rel in fn_rels[1:]:
        doc.part.drop_rel(rel.rId)


def _get_footnotes_xml(doc) -> Optional[object]:
    """获取文档的脚注 XML 根元素."""
    FOOTNOTE_REL = ('http://schemas.openxmlformats.org/officeDocument/'
                    '2006/relationships/footnotes')
    for rel in doc.part.rels.values():
        if rel.reltype == FOOTNOTE_REL:
            from lxml import etree
            return etree.fromstring(rel.target_part.blob)
    return None


def _get_footnotes_part_obj(doc):
    """获取脚注 Part 对象 (用于写回 blob)."""
    FOOTNOTE_REL = ('http://schemas.openxmlformats.org/officeDocument/'
                    '2006/relationships/footnotes')
    for rel in doc.part.rels.values():
        if rel.reltype == FOOTNOTE_REL:
            return rel.target_part
    return None


def format_footnotes(doc, rules: dict = None):
    """格式化现有脚注, 每页重新编号, 数字编号.

    Args:
        doc: Document 对象.
        rules: 自定义规则表 (footnote 元素); None = 默认规则.
    """
    if rules is None:
        from rules import DEFAULT_RULES
        rules = DEFAULT_RULES
    fn_rule = rules['footnote']

    fn_xml = _get_footnotes_xml(doc)
    fn_part_obj = _get_footnotes_part_obj(doc)
    if fn_xml is None:
        return  # 文档无脚注

    # 1) 设置脚注属性: 每页重新编号, 编号格式为 1,2,3...
    fn_pr = fn_xml.find(qn('w:footnotePr'))
    if fn_pr is None:
        fn_pr = parse_xml(f'<w:footnotePr {nsdecls("w")} />')
        fn_xml.insert(0, fn_pr)

    # numFmt: 数字编号
    num_fmt = fn_pr.find(qn('w:numFmt'))
    if num_fmt is None:
        num_fmt = parse_xml(f'<w:numFmt {nsdecls("w")} />')
        fn_pr.append(num_fmt)
    num_fmt.set(qn('w:val'), 'decimal')

    # numStart: 从 1 开始
    num_start = fn_pr.find(qn('w:numStart'))
    if num_start is None:
        num_start = parse_xml(f'<w:numStart {nsdecls("w")} />')
        fn_pr.append(num_start)
    num_start.set(qn('w:val'), '1')

    # numRestart: 0 = 每页重新开始 (restart per page)
    num_restart = fn_pr.find(qn('w:numRestart'))
    if num_restart is None:
        num_restart = parse_xml(f'<w:numRestart {nsdecls("w")} />')
        fn_pr.append(num_restart)
    num_restart.set(qn('w:val'), '0')

    # 2) 逐个脚注段落格式化
    for fn_elem in fn_xml.findall(qn('w:footnote')):
        fn_id_str = fn_elem.get(qn('w:id'))
        if fn_id_str is None:
            continue
        fn_id = int(fn_id_str)
        if fn_id <= 0:
            continue  # 跳过 -1 (分隔线) 和 0 (分隔线续)

        for p_elem in fn_elem.findall(qn('w:p')):
            _format_footnote_paragraph(p_elem, fn_rule)

    # 将修改写回 Part (通过 citation_formatter 的统一写入函数)
    if fn_part_obj is not None:
        try:
            from citation_formatter import _write_footnote_part
            _write_footnote_part(fn_part_obj, fn_xml)
        except ImportError:
            # 降级: 直接使用 _blob
            from lxml import etree as _etree
            fn_part_obj._blob = _etree.tostring(
                fn_xml, xml_declaration=True, encoding='UTF-8', standalone=True)


def _format_footnote_paragraph(p_elem, rule: dict):
    """格式化单个脚注段落 XML 元素.

    Args:
        p_elem: 脚注段落 XML 元素.
        rule: footnote 元素规则 (load_rules 输出).
    """
    # 行距
    pPr = p_elem.find(qn('w:pPr'))
    if pPr is None:
        pPr = parse_xml(f'<w:pPr {nsdecls("w")} />')
        p_elem.insert(0, pPr)

    spacing = pPr.find(qn('w:spacing'))
    if spacing is None:
        spacing = parse_xml(f'<w:spacing {nsdecls("w")} />')
        _insert_pPr_ordered(pPr, spacing)
    ls = rule['line_spacing']
    if isinstance(ls, dict) and ls['mode'] == 'exact':
        spacing.set(qn('w:line'), str(int(ls['value'] * 20)))
        spacing.set(qn('w:lineRule'), 'exact')
    else:
        spacing.set(qn('w:line'), str(int(float(ls) * 240)))
        spacing.set(qn('w:lineRule'), 'auto')
    # beforeLines/afterLines 与 before/after 互斥 (EG_SpaceBefore choice), 先清除
    spacing.attrib.pop(qn('w:beforeLines'), None)
    spacing.attrib.pop(qn('w:afterLines'), None)
    spacing.set(qn('w:before'), '0')
    spacing.set(qn('w:after'), '0')
    # 属性追加顺序违反 CT_Spacing schema → Word 判损坏, 需重排
    _reorder_attrs(spacing, _SPACING_ATTR_ORDER)

    # 每个 run: 按规则 (字号 half-points = 磅值 * 2)
    size_hp = str(int(rule['size'] * 2))
    for r_elem in p_elem.findall(qn('w:r')):
        rPr = r_elem.find(qn('w:rPr'))
        if rPr is None:
            rPr = parse_xml(f'<w:rPr {nsdecls("w")} />')
            r_elem.insert(0, rPr)

        # 字号
        sz = rPr.find(qn('w:sz'))
        if sz is None:
            sz = parse_xml(f'<w:sz {nsdecls("w")} />')
            rPr.append(sz)
        sz.set(qn('w:val'), size_hp)

        szCs = rPr.find(qn('w:szCs'))
        if szCs is None:
            szCs = parse_xml(f'<w:szCs {nsdecls("w")} />')
            rPr.append(szCs)
        szCs.set(qn('w:val'), size_hp)

        # 字体
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} />')
            rPr.insert(0, rFonts)
        rFonts.set(qn('w:eastAsia'), rule['font'])
        rFonts.set(qn('w:ascii'), rule['font'])
        rFonts.set(qn('w:hAnsi'), rule['font'])


def detect_title_range(paragraphs) -> Tuple[int, int]:
    """检测标题段落范围.

    第一个非空段落视为论文题目。若紧随的段落也为非空且非标题,
    则可能是副标题或作者信息, 一并纳入题目区域。
    摘要/关键词属元数据, 不占用作者行上限, 持续并入区域 (修复 BUG-008:
    题目+作者行后摘要/关键词被排除在区域外, 被误判为正文).

    Returns:
        (start_index, end_index_exclusive) — 题目区域在 paragraphs 中的范围.
    """
    start = -1
    end = 0
    extra_count = 0  # 已纳入的非元数据附加段数 (副标题/作者行, 上限 1)

    for i, para in enumerate(paragraphs):
        text = para.text.strip()
        if not text:
            if start >= 0:
                # 空行标志着题目区域结束
                end = i
                break
            continue

        if start < 0:
            start = i
            end = i + 1
            continue

        # 题目之后: 若遇到标题关键词则结束题目区域
        if detect_heading_level(text) > 0:
            end = i
            break

        # 摘要/关键词属元数据: 不占用作者行上限, 持续并入区域
        if is_abstract_or_keywords(text) is not None:
            end = i + 1
            continue

        # 非元数据段落: 可能是副标题/作者行, 最多纳入 1 段
        if extra_count >= 1:
            end = i
            break
        extra_count += 1
        end = i + 1

    return (start, end)


_ALIGN_MAP = {
    'center': WD_ALIGN_PARAGRAPH.CENTER,
    'left': WD_ALIGN_PARAGRAPH.LEFT,
    'right': WD_ALIGN_PARAGRAPH.RIGHT,
    'justify': WD_ALIGN_PARAGRAPH.JUSTIFY,
}


def _apply_para_format(para, rule: dict, body_indent: int = 2):
    """按规则表应用段落格式 (对齐/行距/缩进/字体).

    Args:
        para: 段落对象.
        rule: 元素规则 dict (rules.py 的 load_rules 输出, 字段齐全).
        body_indent: 正文缩进字符数; 仅当 rule['indent'] 为 None 时生效
                     (body 元素默认 None = 由该参数决定).
    """
    set_alignment(para, _ALIGN_MAP[rule['align']])
    set_line_spacing(para, rule['line_spacing'])

    indent = rule['indent']
    if indent is None:
        indent = body_indent
    if indent and indent > 0:
        set_first_line_indent(para, chars=indent)
    else:
        remove_first_line_indent(para)

    for run in para.runs:
        apply_run_font(run, rule['font'], FONT_EN,
                       Pt(rule['size']), bold=rule['bold'])


def format_document(input_path: str, output_path: str,
                    body_indent: int = 2, rules: dict = None,
                    fix_citations: bool = False,
                    citation_rules: dict = None) -> dict:
    """主格式化函数.

    Args:
        input_path: 输入 .docx 文件路径.
        output_path: 输出 .docx 文件路径.
        body_indent: 正文字符缩进数 (默认2, 0=不缩进).
        rules: 自定义格式规则表 (rules.py load_rules 的输出);
               None = 使用默认规则 (skill 预设格式).
        fix_citations: 是否同时修复引注格式;
               True 时一步完成排版+引注修复, 统计写入 stats['citation_fixes'].
        citation_rules: 注释体例 dict (citation_rules.load_citation_rules 输出);
               None = 默认体例.

    Returns:
        统计信息 dict: {title, headings_l1..l4, body, footnotes,
                        citation_fixes?, punct_fixes?, errors}
    """
    if rules is None:
        from rules import DEFAULT_RULES
        rules = DEFAULT_RULES

    stats = {
        'title': 0, 'abstract': 0, 'keywords': 0,
        'headings_l1': 0, 'headings_l2': 0,
        'headings_l3': 0, 'headings_l4': 0,
        'body': 0, 'footnotes': 0, 'errors': 0,
    }

    doc = Document(input_path)
    paragraphs = doc.paragraphs

    # —— 检测题目区域 ——
    title_start, title_end = detect_title_range(paragraphs)
    if title_start < 0:
        stats['errors'] += 1
        print("  ⚠ 未检测到任何文本内容")
        doc.save(output_path)
        return stats

    title_indices = set(range(title_start, title_end))

    # —— 逐段格式化 ——
    for i, para in enumerate(paragraphs):
        text = para.text.strip()
        if not text:
            continue

        # 确保段落有 run
        clear_paragraph_runs(para)

        # 题目区域
        if i in title_indices:
            is_first_title = (i == title_start)
            ak_type = is_abstract_or_keywords(text)

            if is_first_title:
                # 论文题目
                _apply_para_format(para, rules['title'], body_indent)
                stats['title'] += 1
            elif ak_type == 'abstract':
                # 摘要
                _apply_para_format(para, rules['abstract'], body_indent)
                stats.setdefault('abstract', 0)
                stats['abstract'] += 1
            elif ak_type == 'keywords':
                # 关键词
                _apply_para_format(para, rules['keywords'], body_indent)
                stats.setdefault('keywords', 0)
                stats['keywords'] += 1
            else:
                # 其他标题区域内容 (作者名等)
                _apply_para_format(para, rules['author'], body_indent)
            continue

        # 标题 / 正文
        level = detect_heading_level(text)

        if level == 1:
            _apply_para_format(para, rules['h1'], body_indent)
            stats['headings_l1'] += 1
        elif level == 2:
            _apply_para_format(para, rules['h2'], body_indent)
            stats['headings_l2'] += 1
        elif level == 3:
            _apply_para_format(para, rules['h3'], body_indent)
            stats['headings_l3'] += 1
        elif level == 4:
            _apply_para_format(para, rules['h4'], body_indent)
            stats['headings_l4'] += 1
        else:
            # 正文
            _apply_para_format(para, rules['body'], body_indent)
            stats['body'] += 1

    # —— 格式化脚注 ——
    try:
        fn_xml = _get_footnotes_xml(doc)
        if fn_xml is not None:
            fn_count = len([e for e in fn_xml.findall(qn('w:footnote'))
                           if e.get(qn('w:id')) and int(e.get(qn('w:id'))) > 0])
            stats['footnotes'] = fn_count
            format_footnotes(doc, rules=rules)
    except Exception as e:
        print(f"  ⚠ 脚注处理异常: {e}")
        stats['errors'] += 1

    # —— 标点规范化 (正文/摘要等, 跳过题目与各级标题) ——
    try:
        from citation_formatter import normalize_chinese_punctuation
        punct_fixes = 0
        for i, para in enumerate(paragraphs):
            text = para.text.strip()
            if not text:
                continue
            if i == title_start or detect_heading_level(text) > 0:
                continue  # 题目/各级标题: 不修改标点
            for run in para.runs:
                if run.text:
                    new_text = normalize_chinese_punctuation(run.text)
                    if new_text != run.text:
                        punct_fixes += 1
                        run.text = new_text
        if punct_fixes > 0:
            stats['punct_fixes'] = punct_fixes
    except ImportError:
        pass

    # —— 引注格式修复 (可选, 按注释体例) ——
    if fix_citations:
        try:
            from citation_formatter import format_all_footnotes
            cstats = format_all_footnotes(doc, fix=True, style=citation_rules)
            if cstats.get('fixed', 0) > 0:
                stats['citation_fixes'] = cstats['fixed']
        except ImportError:
            pass
        except Exception as e:
            print(f"  ⚠ 引注修复异常: {e}")
            stats['errors'] += 1

    # —— 保存 ——
    # 清理 Word 转换遗留的空 numbering 部件 (否则 Word 判文档损坏)
    _drop_unused_numbering(doc)
    # 去重指向同一 footnotes part 的重复关系 (否则 Word 判文档损坏)
    _dedupe_footnotes_rel(doc)
    doc.save(output_path)
    return stats


def check_document(input_path: str) -> dict:
    """检查模式: 仅读取并报告文档结构, 不修改.

    Returns:
        {paragraphs: 总段数, title: 题目文本, headings: [(层级, 文本),...], ...}
    """
    doc = Document(input_path)
    result = {
        'file': input_path,
        'total_paragraphs': len(doc.paragraphs),
        'title': '',
        'headings': [],
        'has_footnotes': False,
    }

    title_start, title_end = detect_title_range(doc.paragraphs)
    if title_start >= 0:
        titles = [doc.paragraphs[i].text.strip()
                  for i in range(title_start, title_end)
                  if doc.paragraphs[i].text.strip()]
        result['title'] = ' | '.join(titles)

    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text or i < title_end:
            continue
        level = detect_heading_level(text)
        if level > 0:
            result['headings'].append((level, text[:60]))

    result['has_footnotes'] = _get_footnotes_xml(doc) is not None
    return result


def backup_file(filepath: str) -> str:
    """创建备份文件 (原文件名.backup_时间戳.docx)."""
    p = Path(filepath)
    ts = datetime.now().strftime('%Y%m%d_%H%M%S')
    backup = p.parent / f"{p.stem}.backup_{ts}{p.suffix}"
    shutil.copy2(filepath, str(backup))
    return str(backup)


# ---------------------------------------------------------------------------
# CLI
# ---------------------------------------------------------------------------
def build_parser() -> argparse.ArgumentParser:
    """构建命令行解析器 (提取为独立函数, 便于测试默认值)."""
    parser = argparse.ArgumentParser(
        description='中文学术论文格式规范化工具',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
示例:
  python format_paper.py 论文.docx                    # 覆盖原文件 (自动备份)
  python format_paper.py 论文.docx -o 论文_格式.docx  # 输出到新文件
  python format_paper.py 论文.docx --check             # 仅检查不修改
  python format_paper.py 论文/ --batch                 # 批量处理文件夹
  python format_paper.py 论文.docx --body-indent 2     # 正文首行缩进2字符
        """,
    )
    parser.add_argument('input', nargs='?', help='输入 .docx 文件路径, 或文件夹 (需配合 --batch)')
    parser.add_argument('-o', '--output', help='输出文件路径 (默认覆盖原文件)')
    parser.add_argument('--check', action='store_true',
                        help='仅检查文档结构, 不修改')
    parser.add_argument('--batch', action='store_true',
                        help='批量处理文件夹中所有 .docx 文件')
    parser.add_argument('--no-backup', action='store_true',
                        help='覆盖时不自动备份')
    parser.add_argument('--body-indent', type=int, default=2,
                        help='正文首行缩进字符数 (默认2, 0=不缩进)')
    parser.add_argument('--check-prereqs', action='store_true',
                        help='检查依赖和环境')
    parser.add_argument('--diagnostics', action='store_true',
                        help='输出技能元数据')
    parser.add_argument('--fix-citations', action='store_true',
                        help='自动修复脚注中的引注格式 (依据《法学引注手册》)')
    parser.add_argument('--check-citations', action='store_true',
                        help='仅检查脚注引注格式, 不修改')
    parser.add_argument('--rules', default=None,
                        help='自定义格式规则 JSON 配置文件 '
                             '(如 {"title": {"size": "二号", "bold": true}}, 见 references/custom-rules.md)')
    parser.add_argument('--citation-rules', default=None,
                        help='自定义注释体例 JSON 配置文件 '
                             '(见 references/citation-rules.md)')
    parser.add_argument('--preset', default=None,
                        help='预设注释体例: 法学引注手册 / 《法学家》《中外法学》注释体例 / 《中国法学》 / 《法商研究》 / 《法学研究》')

    return parser


def load_rules_or_exit(config_path):
    """加载自定义规则, 出错时打印友好错误并退出 (返回 None = 用默认规则)."""
    if not config_path:
        return None
    try:
        from rules import load_rules
        rules = load_rules(config_path)
        print(f"已加载自定义格式规则: {config_path}")
        return rules
    except (ValueError, FileNotFoundError) as e:
        print(f'{{"error": "{e}", "error_type": "validation", '
              f'"hint": "请检查规则配置文件格式 (见 references/custom-rules.md)"}}',
              file=sys.stderr)
        sys.exit(1)


def load_citation_rules_or_exit(config_path, preset=None):
    """加载自定义注释体例, 出错时打印友好错误并退出 (返回 None = 用默认体例)."""
    if not config_path and not preset:
        return None
    if not _HAS_CITATION_RULES:
        print('{"error": "未找到 citation_rules 模块", '
              '"error_type": "runtime", '
              '"hint": "请确保 citation_rules.py 与 format_paper.py 同目录"}',
              file=sys.stderr)
        sys.exit(1)
    try:
        style = _load_citation_rules(config_path, preset=preset)
        src = config_path or f'预设 {preset}'
        print(f"已加载注释体例: {src}")
        return style
    except (ValueError, FileNotFoundError) as e:
        print(f'{{"error": "{e}", "error_type": "validation", '
              f'"hint": "请检查注释体例配置文件格式 '
              f'(见 references/citation-rules.md)"}}',
              file=sys.stderr)
        sys.exit(1)


def main():
    parser = build_parser()
    args = parser.parse_args()

    # —— 诊断模式 ——
    if args.diagnostics:
        import json
        diag = {
            'skill': 'chinese-paper-format-skill',
            'version': '1.0.0',
            'harness_level': 'production',
            'commands': ['format_paper.py'],
            'harness_features': {
                'input_validation': True,
                'output_sanity': False,
                'check_prereqs': True,
                'diagnostics': True,
            },
        }
        print(json.dumps(diag, ensure_ascii=False, indent=2))
        return

    # —— 依赖检查 ——
    if args.check_prereqs:
        import json as _json
        checks = []

        # Python version
        import platform
        py_ok = sys.version_info >= (3, 8)
        checks.append({
            'check': 'python_version',
            'required': '>=3.8',
            'found': platform.python_version(),
            'ok': py_ok,
        })

        # python-docx
        try:
            import docx as _d
            checks.append({'check': 'python-docx', 'required': '>=1.0.0',
                          'found': _d.__version__, 'ok': True})
        except Exception:
            checks.append({'check': 'python-docx', 'required': '>=1.0.0',
                          'found': 'missing', 'ok': False})

        # lxml
        try:
            import lxml.etree
            checks.append({'check': 'lxml', 'required': '>=4.0',
                          'found': 'available', 'ok': True})
        except Exception:
            checks.append({'check': 'lxml', 'required': '>=4.0',
                          'found': 'missing', 'ok': False})

        ready = all(c['ok'] for c in checks)
        print(_json.dumps({'ready': ready, 'checks': checks},
                          ensure_ascii=False, indent=2))
        sys.exit(0 if ready else 1)

    # —— 输入验证 ——
    if not args.input:
        print('{"error": "请提供输入文件路径", '
              '"error_type": "validation", '
              '"hint": "用法: python format_paper.py <文件.docx>"}',
              file=sys.stderr)
        sys.exit(1)

    input_path = Path(args.input)
    if not input_path.exists():
        print(f'{{"error": "文件不存在: {args.input}", '
              f'"error_type": "validation", '
              f'"hint": "请检查文件路径是否正确"}}',
              file=sys.stderr)
        sys.exit(1)

    # —— 加载自定义格式规则 ——
    rules = load_rules_or_exit(args.rules)

    # —— 加载自定义注释体例 ——
    citation_rules = load_citation_rules_or_exit(args.citation_rules,
                                                  preset=args.preset)

    # —— 批量模式 ——
    if args.batch:
        if not input_path.is_dir():
            print(f'{{"error": "--batch 需要输入文件夹", '
                  f'"error_type": "validation"}}',
                  file=sys.stderr)
            sys.exit(1)

        docx_files = list(input_path.glob('*.docx'))
        if not docx_files:
            print(f'{{"error": "文件夹中未找到 .docx 文件", '
                  f'"error_type": "validation"}}',
                  file=sys.stderr)
            sys.exit(1)

        print(f"批量处理 {len(docx_files)} 个文件...\n")
        total_stats = {'files': len(docx_files), 'ok': 0, 'fail': 0}

        for f in docx_files:
            try:
                bkp = backup_file(str(f))
                print(f"  {f.name}  (备份: {Path(bkp).name})")
                stats = format_document(str(f), str(f),
                                        body_indent=args.body_indent,
                                        rules=rules,
                                        fix_citations=args.fix_citations,
                                        citation_rules=citation_rules)
                _print_stats(stats)
                if stats.get('citation_fixes', 0) > 0:
                    print(f"    引注修复: {stats['citation_fixes']} 处")

                total_stats['ok'] += 1
            except Exception as e:
                print(f"  ✗ {f.name}: {e}")
                total_stats['fail'] += 1

        print(f"\n完成: {total_stats['ok']} 成功, {total_stats['fail']} 失败")
        sys.exit(0 if total_stats['fail'] == 0 else 1)

    # —— 格式检查 (只读) ——
    if args.check:
        result = check_document(str(input_path))
        print(f"文件: {result['file']}")
        print(f"总段落: {result['total_paragraphs']}")
        print(f"题目: {result['title'][:80] if result['title'] else '(未检测到)'}")
        print(f"标题 ({len(result['headings'])} 个):")
        for level, txt in result['headings']:
            prefix = {1: '一、', 2: '（一）', 3: '1. ', 4: '（1）'}.get(level, '?')
            print(f"  L{level} [{prefix}] {txt}")
        print(f"脚注: {'有' if result['has_footnotes'] else '无'}")
        return

    # —— 单文件格式化 ——
    # .doc 文件警告
    if input_path.suffix.lower() == '.doc':
        print("⚠ .doc 格式不支持直接处理。请先用 Word 或 LibreOffice 另存为 .docx。")
        print("  Word: 文件 → 另存为 → Word 文档 (*.docx)")
        print("  LibreOffice: soffice --headless --convert-to docx 文件.doc")
        sys.exit(1)

    if input_path.suffix.lower() not in ('.docx',):
        print(f'{{"error": "不支持的文件格式: {input_path.suffix}", '
              f'"error_type": "validation", '
              f'"hint": "请提供 .docx 文件"}}',
              file=sys.stderr)
        sys.exit(1)

    # 确定输出路径
    output_path = args.output if args.output else str(input_path)

    # 覆盖时自动备份
    if not args.output and not args.no_backup:
        bkp = backup_file(str(input_path))
        print(f"已备份: {Path(bkp).name}")

    # 执行格式化
    try:
        print(f"格式化: {input_path}")
        stats = format_document(str(input_path), output_path,
                                body_indent=args.body_indent,
                                rules=rules,
                                fix_citations=args.fix_citations,
                                citation_rules=citation_rules)
        print(f"输出: {output_path}")
        _print_stats(stats)
        if stats.get('citation_fixes', 0) > 0:
            print(f"    引注修复: {stats['citation_fixes']} 处")

        # —— 引注格式检查 (只读) ——
        if args.check_citations:
            if not _HAS_CITATION:
                print("\n⚠ 引注格式化模块未找到, 请确保 citation_formatter.py 在同一目录")
            else:
                _run_citation_checks(output_path, fix=False, style=citation_rules)

        print("✓ 完成")
    except Exception as e:
        print(f'{{"error": "{e}", "error_type": "runtime", '
              f'"hint": "格式化过程中出现错误, 请检查文档是否损坏"}}',
              file=sys.stderr)
        sys.exit(1)


def _run_citation_checks(filepath: str, fix: bool = False, style=None):
    """运行引注格式检查/修复."""
    doc = Document(filepath)
    if fix:
        print("\n🔧 修复引注格式...")
        cstats = format_all_footnotes(doc, fix=True, style=style)
        doc.save(filepath)
        print(f"  检查脚注: {cstats['total']} 条")
        print(f"  自动修复: {cstats['fixed']} 处")
        print(f"  剩余问题: {cstats['issues']} 个")
        if cstats['issues'] > 0:
            print(f"  ⚠ {cstats['issues']} 个问题需手动处理")
    else:
        print("\n📋 检查引注格式...")
        footnotes = extract_footnotes(doc)
        total_issues = 0
        for fn in footnotes:
            issues = check_footnote(fn['full_text'])
            total_issues += len(issues)
            if issues:
                print(f"  [脚注 {fn['id']}] {fn['full_text'][:60]}...")
                for iss in issues:
                    sev = {'high': '🔴', 'medium': '🟡', 'low': '🟢', 'info': 'ℹ️'}
                    print(f"    {sev.get(iss['severity'], '  ')} {iss['message']}")
        if total_issues == 0:
            print("  ✓ 引注格式正常")
        else:
            print(f"  共 {total_issues} 个问题，运行 --fix-citations 自动修复")


def _print_stats(stats: dict):
    """打印格式化统计."""
    parts = []
    if stats.get('title'):
        parts.append(f"题目: {stats['title']}")
    if stats.get('abstract'):
        parts.append(f"摘要: {stats['abstract']}")
    if stats.get('keywords'):
        parts.append(f"关键词: {stats['keywords']}")
    if stats.get('headings_l1'):
        parts.append(f"一级标题: {stats['headings_l1']}")
    if stats.get('headings_l2'):
        parts.append(f"二级标题: {stats['headings_l2']}")
    if stats.get('headings_l3'):
        parts.append(f"三级标题: {stats['headings_l3']}")
    if stats.get('body'):
        parts.append(f"正文段落: {stats['body']}")
    if stats.get('footnotes'):
        parts.append(f"脚注: {stats['footnotes']}")
    if stats.get('punct_fixes'):
        parts.append(f"标点规范: {stats['punct_fixes']}处")
    if stats.get('errors'):
        parts.append(f"异常: {stats['errors']}")
    if parts:
        print(f"  {' | '.join(parts)}")


if __name__ == '__main__':
    main()
