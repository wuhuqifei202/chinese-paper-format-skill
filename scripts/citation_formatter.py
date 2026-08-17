#!/usr/bin/env python3
"""
Citation Formatter for Chinese Legal Academic Papers
法学引注格式规范化模块

依据《法学引注手册》(2019) 的规范，检测并修复脚注中的引注格式问题。

支持的格式规则:
  - 中文标点符号规范化: ，、：。《》 vs ,.:《》"
  - 作者与文献名称之间用冒号
  - 文献名称使用书名号
  - 期刊/报纸/文集文章前加"载"
  - 页码格式: 第×页 / 第×-×页
  - 同一注释多文献用分号分隔
  - 引用符号位置 (句号后 vs 句中)
  - 同一文献多次出现时的略写格式
  - 引领词规范: 参见/见/又见/另见/转引自

Usage:
    python citation_formatter.py input.docx --check          # 仅检查
    python citation_formatter.py input.docx --fix            # 自动修复
    python citation_formatter.py input.docx -o output.docx   # 输出到新文件
"""

import re
import sys
import os
import argparse
from pathlib import Path
from typing import List, Dict, Optional, Tuple

from docx import Document
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# 确保脚本目录在 sys.path 中 — 裸 import citation_rules 依赖于此
# (以 `python scripts/xxx.py` 运行时 sys.path[0] 恰为脚本目录; 任意 cwd/入口下需显式注入)
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

# 注释体例模块 (可选导入: 提供默认体例 + 自定义配置加载)
try:
    from citation_rules import load_citation_rules, describe_style
    _HAS_CITATION_RULES = True
except ImportError:
    _HAS_CITATION_RULES = False

# ---------------------------------------------------------------------------
# 中文标点映射 — 用于修复英文标点
# ---------------------------------------------------------------------------
# 注意: 不是简单替换，需要在语境中判断
_EN_TO_CN_PUNCT = {
    # 英文 → 中文 (仅当在中文语境中时替换)
    ',': '，',
    ':': '：',
    ';': '；',
    '(': '（',
    ')': '）',
}

# 需要保护的标点 (在数字、URL、外文等上下文中不替换)
_PROTECTED_CONTEXTS = [
    r'\d\.\d',           # 数字中的点号 (如 3.14)
    r'https?://',         # URL
    r'[A-Za-z]\.',        # 英文缩写
    r'\d,\d',             # 数字中的逗号
]

# ---------------------------------------------------------------------------
# 引注格式正则模式 — 用于检测和验证
# ---------------------------------------------------------------------------

# 期刊文章: 作者：《标题》，载《期刊》年份年第×期
_RE_JOURNAL = re.compile(
    r'([^，。：]+)[：:]\s*《([^》]+)》\s*[,，]\s*载《([^》]+)》\s*'
    r'(\d{4})\s*年\s*第\s*(\d+)\s*期'
)

# 图书: 作者：《书名》，出版社年份版
_RE_BOOK = re.compile(
    r'([^，。：]+)[：:]\s*《([^》]+)》\s*[,，]\s*'
    r'(.+出版社)\s*(\d{4})\s*年版'
)

# 报纸: 作者：《标题》，载《报纸》年月日
_RE_NEWSPAPER = re.compile(
    r'([^，。：]+)[：:]\s*《([^》]+)》\s*[,，]\s*载《([^》]+)》\s*'
    r'(\d{4})\s*年\s*(\d{1,2})\s*月\s*(\d{1,2})\s*日'
)

# 文集文章: 作者：《标题》，载编者编：《文集》，出版社年份版
_RE_BOOK_ARTICLE = re.compile(
    r'([^，。：]+)[：:]\s*《([^》]+)》\s*[,，]\s*'
    r'载\s*([^：:]+)[：:]\s*《([^》]+)》'
)

# 翻译作品: [国籍]作者：《书名》，译者译，出版社年份版
_RE_TRANSLATION = re.compile(
    r'\[([^\]]+)\]\s*([^：:]+)[：:]\s*《([^》]+)》\s*[,，]\s*'
    r'([^，]+译)\s*[,，]\s*(.+)'
)

# 页码: 第×页 或 第×-×页
_RE_PAGE = re.compile(r'第\s*(\d+(?:[-–—]\d+)?)\s*页')

# 引领词
_LEADING_WORDS = ['参见', '见', '又见', '另见', '转引自', '同上注', '同前注']

# —— 法律文件引用检测 ——

# 规范性文件: 制定机关《名称》，文号
_RE_NORMATIVE_DOC = re.compile(
    r'([一-鿿]+(?:[部委局办院署会])?)\s*《([^》]+)》\s*[,，]\s*'
    r'([一-鿿A-Za-z]+[〔\[发]?\s*[〔\[]?\d{4}[〕\]]?\s*\]?\d+\s*号)'
)

# 文号中的方括号错误: 国发[2007]19号 → 应为六角括号
_RE_WRONG_BRACKET = re.compile(
    r'([一-鿿A-Za-z]+[发字]?)\s*\[\s*(\d{4})\s*\]\s*(\d+\s*号)'
)

# 文号中缺六角括号: 国发2007 19号
_RE_MISSING_BRACKET = re.compile(
    r'([一-鿿A-Za-z]+[发字]?)\s*(\d{4})\s+(\d+\s*号)'
)

# 法律条文引用: 《法律名称》第×条
_RE_LEGAL_ARTICLE = re.compile(
    r'《([^》]+)》\s*第\s*\d+\s*条'
)

# 法规引用: 制定机关《法规名称》（年份）
_RE_REGULATION = re.compile(
    r'([一-鿿]+(?:[部委局办院署会]))\s*《([^》]+)》\s*[（(]\s*(\d{4})\s*[）)]'
)

# —— 案例引用检测 ——

# 案例名称: ×××诉××××××案
_RE_CASE_NAME = re.compile(
    r'([一-鿿A-Za-z（）()]+(?:诉|与)\s*[一-鿿A-Za-z（）()]+[\s\w]*?案)'
)

# 案号: 法院名称文书名称，（年份）案号
_RE_CASE_NUMBER = re.compile(
    r'[（(]\s*(\d{4})\s*[）)]\s*[一-鿿A-Za-z0-9]+(?:字第?)?\d+\s*号'
)

# 指导案例: 最高人民法院指导案例×号（年份）
_RE_GUIDING_CASE = re.compile(
    r'(?:最高人民法院|最高人民检察院)\s*指导[性案]例\s*\d+\s*号'
)

# 案号中方括号错误: [1998]海行初字第142号 → (1998)
_RE_CASE_BRACKET = re.compile(
    r'[\[【]\s*(\d{4})\s*[\]】]\s*([一-鿿A-Za-z0-9]+(?:字第?)?\d+\s*号)'
)

# "字第" 冗余: （1998）海行初字第142号 → 规范写法保留"字"在旧案号
# 旧案号保留"字", 新案号不写"字" (已由最高人民法院规范)
# 检测旧案号中的"字" (信息性提示)
_RE_OLD_CASE_NUMBER = re.compile(r'字第\d+\s*号')

# 常见错误模式
# ---------------------------------------------------------------------------
# 自动修复规则 (按顺序应用)
# ---------------------------------------------------------------------------
# 规则按 结构性修复 → 标点修复 → 空格清理 的顺序排列。
# 顺序原则:
#   1. 先修结构 (作者冒号、页码、出版社)，再修标点，最后清空格。
#   2. 法律/案例规则 (7-9) 不与文献规则 (1-5) 重叠匹配域。
#   3. 写回 _ERROR_PATTERNS 前，新规则必须在本测试文件中验证其位置不
#      与已有序号冲突。相关测试: tests/test_citation_rules.py::TestPatternOrder
# ---------------------------------------------------------------------------
# ── 中文标点规范化 (公共函数, 供正文和脚注共用) ──

_PUNCTUATION_PATTERNS = [
    # 句中: 中文后英文标点再接中文
    (re.compile(r'([一-鿿）》）])\.(?=[\s一-鿿（《])'), r'\1。'),
    (re.compile(r'([一-鿿）》》]),(?=\s*[一-鿿（《])'), r'\1，'),
    (re.compile(r'([一-鿿）》》]);(?=\s*[一-鿿（《])'), r'\1；'),
    (re.compile(r'([一-鿿）》》]):(?=\s*[一-鿿（《])'), r'\1：'),
    # 英文/数字后、中文前的英文逗号 (如 URL 后接中文句, 保护 1,000)
    (re.compile(r'(?<=[A-Za-z0-9]),(?=\s*[一-鿿（《])'), r'，'),
    # 英文括号包中文 (含 e.g./i.e. 等缩写混排, 整对转换; 保护纯数字/纯英文)
    (re.compile(r'\(([^()]*[一-鿿][^()]*)\)'), r'（\1）'),
    # 终端句号 (在括号成对转换之后执行, 确保全角 ）后句号也能转)
    (re.compile(r'([一-鿿》）》\d])\.[\s]*$'), r'\1。'),
    # 中文问号 / 叹号
    (re.compile(r'([一-鿿）》》])\?'), r'\1？'),
    (re.compile(r'([一-鿿）》》])!'), r'\1！'),
    # 中文后独立英文括号 (兜底: 修复输入中已混排的半角括号)
    (re.compile(r'([一-鿿》）])\((?=[一-鿿（《])'), r'\1（'),
    (re.compile(r'(?<=[一-鿿》）])\)'), r'）'),
    # 空格清理
    (re.compile(r'([，。：；、）》》）])\s+'), r'\1'),
    (re.compile(r'\s+([，。：；、《（])'), r'\1'),
]


def normalize_chinese_punctuation(text: str) -> str:
    """将中文上下文中的英文标点统一替换为中文标点.

    同时清理多余空格。不修改数字、案号、文号中的标点。
    可供正文和脚注共同调用。
    """
    for pattern, replacement in _PUNCTUATION_PATTERNS:
        text = pattern.sub(replacement, text)
    return text


# ── 引注自动修复规则 ──

# 修复规则按「组」组织, 便于按注释体例开关 (见 citation_rules.py 的 rules)。
# _ERROR_PATTERNS 为全部规则的扁平拼接, 顺序保持不变 (结构性修复 → 标点 → 空格),
# 供既有测试与调用方使用。

# 1. 作者名后逗号+书名号 → 冒号+书名号 (如 "王利明, 《" → "王利明：《")
_RULES_AUTHOR_SEPARATOR = [
    (re.compile(r'([一-鿿]{2,4})\s*[,，]\s*(《[^》]+》)'), r'\1：\2'),
]

# 2. 英文页码格式 p.xx → 第xx页
_RULES_PAGE_FORMAT = [
    (re.compile(r'[,，]\s*[Pp]\.\s*(\d+)'), r'，第\1页'),
    (re.compile(r'([。；])\s*[Pp]\.\s*(\d+)'), r'\1第\2页'),
]

# 3. 中英文标点混用 (中文上下文中统一使用中文标点)
_RULES_PUNCTUATION = [
    (re.compile(r'([一-鿿）》）])\.(?=[\s一-鿿（《])'), r'\1。'),
    (re.compile(r'([一-鿿）》》]),(?=\s*[一-鿿（《])'), r'\1，'),
    (re.compile(r'([一-鿿）》》]);(?=\s*[一-鿿（《])'), r'\1；'),
    (re.compile(r'([一-鿿）》》]):(?=\s*[一-鿿（《])'), r'\1：'),
    (re.compile(r'([一-鿿》）》\d])\.[\s]*$'), r'\1。'),
    (re.compile(r'([一-鿿）》》])\?'), r'\1？'),
    (re.compile(r'([一-鿿）》》])!'), r'\1！'),
    (re.compile(r'\(([一-鿿][^)]*[一-鿿》])\)'), r'（\1）'),
    (re.compile(r'([一-鿿》）])\((?=[一-鿿（《])'), r'\1（'),
    (re.compile(r'(?<=[一-鿿》）])\)'), r'）'),
]

# 4. "载"字缺失 (期刊/报纸文章 — 文章名后逗号+期刊名+年份)
_RULES_ZAI = [
    (re.compile(r'，《([^》]+)》(\d{4})\s*年\s*第'), r'，载《\1》\2年第'),
]

# 5. 出版社格式不完整 (缺"年版")
_RULES_PUBLISHER = [
    (re.compile(r'([一-鿿]+出版社)\s*(\d{4})\s*年\s*[,，]'), r'\1\2年版，'),
    (re.compile(r'([一-鿿]+出版社)\s*(\d{4})\s*年\s*([。；])'), r'\1\2年版\3'),
]

# 6. 文号方括号 → 六角括号: 国发[2007]19号 → 国发〔2007〕19号
_RULES_LEGAL_BRACKET = [
    (re.compile(r'(国发|法释|法发|法〔[^〕]+〕|[一-鿿]+[发字])\s*\[\s*(\d{4})\s*\]\s*(\d+\s*号)'),
     r'\1〔\2〕\3'),
]

# 7. 案号年份方括号 → 圆括号: [1998]→(1998); 及缺少年份括号
_RULES_CASE_BRACKET = [
    (re.compile(r'[\[【]\s*(\d{4})\s*[\]】]\s*([一-鿿A-Za-z].+?号)'),
     r'（\1）\2'),
    (re.compile(r'(?<=\s)(\d{4})\s+([一-鿿][一-鿿A-Za-z]+(?:字第?)?\d+\s*号)'),
     r'（\1）\2'),
]

# 8. 多余空格 (最后执行, 清除标点替换后的残留空格)
_RULES_SPACES = [
    (re.compile(r'([，。：；、）》》）])\s+'), r'\1'),
    (re.compile(r'\s+([，。：；、《（])'), r'\1'),
]

# —— 默认注释体例的扩展修复规则 (受体例开关控制, 不并入 _ERROR_PATTERNS) ——

# 版次归位: 《书名》第N版 → 《书名》（第N版）
_RULES_EDITION = [
    (re.compile(r'《([^》]+)》(第\s*\d+\s*版)'), r'《\1》（\2）'),
]

# 译者署名: 译者名与 "译" 之间的空格清理 ("黄家镇 译" → "黄家镇译")
_RULES_TRANSLATOR = [
    (re.compile(r'([一-鿿]{2,4}(?:、[一-鿿]{2,4})*)\s+译(?=[，。；）]|$)'), r'\1译'),
]

# 扁平规则表 (保持原顺序, 供既有测试/调用使用)
_ERROR_PATTERNS = (
    _RULES_AUTHOR_SEPARATOR +
    _RULES_PAGE_FORMAT +
    _RULES_PUNCTUATION +
    _RULES_ZAI +
    _RULES_PUBLISHER +
    _RULES_LEGAL_BRACKET +
    _RULES_CASE_BRACKET +
    _RULES_SPACES
)

# 规则组名 → 规则组列表 (对应 citation_rules.py 的 rules 开关键名)
_RULE_GROUPS = {
    'author_separator': _RULES_AUTHOR_SEPARATOR,
    'page_format': _RULES_PAGE_FORMAT,
    'punctuation': _RULES_PUNCTUATION,
    'zai_marker': _RULES_ZAI,
    'publisher_year': _RULES_PUBLISHER,
    'legal_bracket': _RULES_LEGAL_BRACKET,
    'case_bracket': _RULES_CASE_BRACKET,
    'edition': _RULES_EDITION,
    'translator': _RULES_TRANSLATOR,
}


def _resolve_style(style):
    """解析体例: None 时加载默认体例 (若无 citation_rules 模块则返回 None)."""
    if style is not None:
        return style
    if _HAS_CITATION_RULES:
        return load_citation_rules()
    return None


def _rule_enabled(style, name: str) -> bool:
    """判断某修复规则是否启用 (style 为 None 或未配置时默认启用)."""
    if style is None or 'rules' not in style:
        return True
    return style['rules'].get(name, True)


def _patterns_for_style(style) -> List[Tuple]:
    """按体例的 rules 开关返回要应用的 (pattern, replacement) 列表."""
    if not _HAS_CITATION_RULES or style is None:
        return list(_ERROR_PATTERNS)
    patterns: List[Tuple] = []
    for name, group in _RULE_GROUPS.items():
        if not _rule_enabled(style, name):
            continue
        patterns.extend(group)
    patterns.extend(_RULES_SPACES)  # 空格清理总是执行
    return patterns


def _get_footnotes_xml(doc) -> Optional[object]:
    """获取文档的脚注 XML 根元素."""
    FOOTNOTE_REL = ('http://schemas.openxmlformats.org/officeDocument/'
                    '2006/relationships/footnotes')
    for rel in doc.part.rels.values():
        if rel.reltype == FOOTNOTE_REL:
            from lxml import etree
            return etree.fromstring(rel.target_part.blob)
    return None


def extract_footnotes(doc: Document) -> List[Dict]:
    """从文档中提取所有脚注文本.

    Returns:
        [{id, paragraphs: [text], full_text: str}, ...]
    """
    fn_xml = _get_footnotes_xml(doc)
    if fn_xml is None:
        return []

    footnotes = []

    for fn_elem in fn_xml.findall(qn('w:footnote')):
        fn_id = int(fn_elem.get(qn('w:id'), '-1'))
        if fn_id <= 0:
            continue  # 跳过分隔脚注

        paragraphs = []
        for p_elem in fn_elem.findall(qn('w:p')):
            texts = []
            for r_elem in p_elem.findall(qn('w:r')):
                for t_elem in r_elem.findall(qn('w:t')):
                    if t_elem.text:
                        texts.append(t_elem.text)
            para_text = ''.join(texts)
            if para_text.strip():
                paragraphs.append(para_text)

        full_text = '\n'.join(paragraphs)
        footnotes.append({
            'id': fn_id,
            'paragraphs': paragraphs,
            'full_text': full_text.strip(),
        })

    return footnotes


def check_footnote(fn_text: str) -> List[Dict]:
    """检查单条脚注文本的格式问题.

    Returns:
        [{type, severity, message, suggestion, position}, ...]
    """
    issues = []

    # 1. 检查中英文标点混用
    _check_punctuation_mixing(fn_text, issues)

    # 2. 检查书名号使用
    _check_book_title_marks(fn_text, issues)

    # 3. 检查"载"字使用
    _check_zai_usage(fn_text, issues)

    # 4. 检查作者-标题分隔符
    _check_author_title_separator(fn_text, issues)

    # 5. 检查页码格式
    _check_page_format(fn_text, issues)

    # 6. 检查引注符号位置 (在正文中检查, 这里跳过)

    # 7. 检查多余空格
    _check_extra_spaces(fn_text, issues)

    # 8. 检查引领词
    _check_leading_words(fn_text, issues)

    # 9. 检查出版社格式
    _check_publisher_format(fn_text, issues)

    # 10. 检查法律文件引用格式
    _check_legal_doc_format(fn_text, issues)

    # 11. 检查案例引用格式
    _check_case_citation(fn_text, issues)

    return issues


def _check_punctuation_mixing(text: str, issues: List[Dict]):
    """检测中文上下文中使用英文标点的问题."""
    # 中文句号后跟英文逗号
    if re.search(r'[一-鿿]\)\s*,', text):
        issues.append({
            'type': 'punctuation',
            'severity': 'medium',
            'message': '中文上下文中使用了英文标点',
            'suggestion': '将英文标点替换为中文标点（，、：；。）',
            'context': text[:80],
        })

    # 检查常见英文标点在中文后出现
    en_punct_in_cn = []
    for en_p, cn_p in _EN_TO_CN_PUNCT.items():
        if re.search(rf'[一-鿿》）]{re.escape(en_p)}\s', text):
            en_punct_in_cn.append(f'{en_p}→{cn_p}')

    if en_punct_in_cn:
        issues.append({
            'type': 'punctuation',
            'severity': 'low',
            'message': f'中文后使用了英文标点: {", ".join(en_punct_in_cn)}',
            'suggestion': '自动替换为中文标点',
            'context': text[:80],
        })


def _check_book_title_marks(text: str, issues: List[Dict]):
    """检测书名号使用是否规范."""
    # 文献标题应该用书名号
    # 检测"载"后面跟的内容是否用了书名号
    zai_match = re.search(r'载\s*([^《\s][^，。]+?)(?:\s*\d{4})', text)
    if zai_match:
        source_name = zai_match.group(1)
        if '《' not in source_name[:10]:
            issues.append({
                'type': 'title_mark',
                'severity': 'high',
                'message': f'"载"后的来源名称未使用书名号: {source_name[:30]}',
                'suggestion': f'应为: 载《{source_name.strip()}》',
                'context': text[:80],
            })

    # 检测期刊名是否用了书名号 (常见模式: "载XXX年第" 缺少书名号)
    journal_no_mark = re.search(r'载\s*([A-Za-z一-鿿（）]+?)\s*(\d{4})\s*年\s*第\s*\d+\s*期', text)
    if journal_no_mark and '《' not in journal_no_mark.group(1):
        jname = journal_no_mark.group(1)
        if jname.strip():
            issues.append({
                'type': 'title_mark',
                'severity': 'high',
                'message': f'期刊名缺少书名号: {jname}',
                'suggestion': f'应为: 《{jname}》',
                'context': text[:80],
            })


def _check_zai_usage(text: str, issues: List[Dict]):
    """检测"载"字是否按规定使用."""
    # 期刊文章: 应有"载《期刊名》"
    has_journal = re.search(r'\d{4}\s*年\s*第\s*\d+\s*期', text)
    if has_journal and '载' not in text:
        issues.append({
            'type': 'zai_missing',
            'severity': 'high',
            'message': '期刊文章缺少"载"字引导',
            'suggestion': '在文章名后、期刊名前添加"载"字',
            'context': text[:80],
        })

    # 报纸文章: 应有"载《报纸名》年月日"
    has_newspaper = re.search(r'\d{4}\s*年\s*\d{1,2}\s*月\s*\d{1,2}\s*日', text)
    if has_newspaper and '载' not in text and not re.search(r'出版社', text):
        issues.append({
            'type': 'zai_missing',
            'severity': 'high',
            'message': '报纸/新闻类文章缺少"载"字引导',
            'suggestion': '在文章名后、报纸名前添加"载"字',
            'context': text[:80],
        })


def _check_author_title_separator(text: str, issues: List[Dict]):
    """检测作者与文献名称之间的分隔符."""
    # 作者名后应该用冒号（：）而非其他符号分隔
    # 检测: 中文字符+英文冒号+书名号 → 应改为中文冒号
    wrong_sep = re.search(r'([一-鿿）\]\)])[:：]\s*(?![0-9])', text)
    # 这个太宽泛了，聚焦在作者-书名模式
    # 检测 姓名字段后跟书名号但用逗号分隔的情况
    comma_before_title = re.search(r'([一-鿿]{2,4}),\s*《', text)
    if comma_before_title:
        issues.append({
            'type': 'separator',
            'severity': 'medium',
            'message': '作者与文献名称之间使用了逗号而非冒号',
            'suggestion': f'将",《" 改为 "：《"',
            'context': comma_before_title.group()[:40],
        })


def _check_page_format(text: str, issues: List[Dict]):
    """检查页码格式."""
    # 应使用"第×页"而非"p.×"或"第×"
    if re.search(r'[Pp]\.\s*\d+', text):
        issues.append({
            'type': 'page_format',
            'severity': 'medium',
            'message': '使用了英文页码格式 (p.×)',
            'suggestion': '改为中文格式: 第×页',
            'context': text[:80],
        })

    # 页码范围应使用短横线 "-" 而非 "~" 或 "到"
    if re.search(r'第\s*\d+\s*[~～到至]\s*\d+\s*页', text):
        issues.append({
            'type': 'page_format',
            'severity': 'low',
            'message': '页码范围连接符不规范',
            'suggestion': '使用短横线: 第×-×页',
            'context': text[:80],
        })


def _check_extra_spaces(text: str, issues: List[Dict]):
    """检测不必要的空格."""
    # 中文标点前后不应有空格
    extra_space = re.findall(r'[一-鿿]\s{2,}[一-鿿]', text)
    if extra_space:
        issues.append({
            'type': 'spacing',
            'severity': 'low',
            'message': f'存在多余空格 ({len(extra_space)} 处)',
            'suggestion': '删除中文标点前后多余的空格',
            'context': text[:80],
        })


def _check_leading_words(text: str, issues: List[Dict]):
    """检查引领词使用是否规范."""
    # "转引自" 应有原始来源和转引来源
    if '转引自' in text:
        if not re.search(r'[。；]', text):
            issues.append({
                'type': 'leading_word',
                'severity': 'info',
                'message': '使用"转引自"但未标注转引来源',
                'suggestion': '应同时标注原始出处和转引出处',
                'context': text[:80],
            })

    # "参见" vs "见" 的用法提醒
    if '参见' in text and re.search(r'["""]', text):
        issues.append({
            'type': 'leading_word',
            'severity': 'info',
            'message': '直接引用原文建议使用"见"而非"参见"',
            'suggestion': '概括引用用"参见"，直接引用用"见"',
            'context': text[:80],
        })


def _check_publisher_format(text: str, issues: List[Dict]):
    """检查出版社信息格式."""
    # 出版社前不应写城市名
    city_pub = re.search(r'([一-鿿]{2,3})[：:]\s*([一-鿿]+出版社)', text)
    if city_pub:
        issues.append({
            'type': 'publisher',
            'severity': 'low',
            'message': f'出版社前有多余的城市名: {city_pub.group(1)}',
            'suggestion': f'只写"{city_pub.group(2)}"，不写城市名',
            'context': text[:80],
        })


def _check_legal_doc_format(text: str, issues: List[Dict]):
    """检测法律文件引用格式问题."""
    # 1. 文号方括号 → 六角括号: 国发[2007]19号 → 国发〔2007〕19号
    bracket_matches = _RE_WRONG_BRACKET.findall(text)
    for m in bracket_matches:
        wrong = f'{m[0]}[{m[1]}]{m[2]}'
        correct = f'{m[0]}〔{m[1]}〕{m[2]}'
        issues.append({
            'type': 'legal_bracket',
            'severity': 'high',
            'message': f'文号使用了方括号而非六角括号: {wrong}',
            'suggestion': f'应为: {correct}',
            'context': text[:80],
        })

    # 2. 文号缺少括号: 国发2007 19号 → 国发〔2007〕19号
    if not bracket_matches:
        missing = _RE_MISSING_BRACKET.findall(text)
        for m in missing:
            wrong = f'{m[0]}{m[1]} {m[2]}'
            correct = f'{m[0]}〔{m[1]}〕{m[2]}'
            issues.append({
                'type': 'legal_bracket',
                'severity': 'high',
                'message': f'文号缺少年份括号: {wrong}',
                'suggestion': f'应为: {correct}',
                'context': text[:80],
            })

    # 3. 法规缺少制定机关和年份
    law_match = _RE_LEGAL_ARTICLE.search(text)
    if law_match:
        law_name = law_match.group(1)
        # 检测是否为法规（非法律）但未写制定机关
        reg_keywords = ['条例', '办法', '规定', '细则', '通知', '意见']
        if any(kw in law_name for kw in reg_keywords):
            if not _RE_REGULATION.search(text):
                issues.append({
                    'type': 'legal_doc',
                    'severity': 'info',
                    'message': f'法规/规章引用建议标明制定机关和年份: {law_name}',
                    'suggestion': '参照手册第63条格式: 制定机关《法规名称》（年份）',
                    'context': text[:80],
                })


def _check_case_citation(text: str, issues: List[Dict]):
    """检测案例引用格式问题."""
    # 1. 案号年份方括号 → 圆括号: [1998]海行初字第142号 → (1998)
    case_brackets = _RE_CASE_BRACKET.findall(text)
    for year, rest in case_brackets:
        wrong = f'[{year}]{rest}'
        correct = f'（{year}）{rest}'
        issues.append({
            'type': 'case_bracket',
            'severity': 'high',
            'message': f'案号年份应使用圆括号而非方括号: {wrong}',
            'suggestion': f'应为: {correct}',
            'context': text[:80],
        })

    # 2. 指导案例格式检查
    guiding = _RE_GUIDING_CASE.search(text)
    if guiding:
        # 检查是否有发布年份括号
        if not re.search(r'指导[性案]例\s*\d+\s*号\s*[（(]\s*\d{4}\s*[）)]', text):
            issues.append({
                'type': 'case_format',
                'severity': 'medium',
                'message': '指导案例建议用括号标注发布年份',
                'suggestion': '格式: 最高人民法院指导案例24号（2014年）',
                'context': text[:80],
            })

    # 3. 案例名称缺少"案"字
    case_match = _RE_CASE_NAME.search(text)
    if case_match:
        case_name = case_match.group(1)
        if not case_name.endswith('案'):
            issues.append({
                'type': 'case_format',
                'severity': 'low',
                'message': f'案例名称建议以"案"字结尾: {case_name[:40]}',
                'suggestion': '民事/行政案例格式: ×××诉××××××案',
                'context': text[:80],
            })

    # 4. 《最高人民法院公报》案例
    gazette = re.search(r'《最高人民法院公报》', text)
    if gazette:
        # 应有年份和期号
        if not re.search(r'《最高人民法院公报》\s*\d{4}\s*年\s*第\s*\d+\s*期', text):
            issues.append({
                'type': 'case_format',
                'severity': 'medium',
                'message': '《最高人民法院公报》案例缺少年份和期号',
                'suggestion': '格式: 《最高人民法院公报》2015年第11期',
                'context': text[:80],
            })


def auto_fix_footnote(text: str, style=None) -> Tuple[str, int]:
    """自动修复脚注文本中的常见格式问题.

    Args:
        text: 脚注文本.
        style: 注释体例 dict (citation_rules.load_citation_rules 输出);
               None = 默认体例. 体例的 rules 开关决定启用哪些修复规则.

    Returns:
        (fixed_text, fix_count) — 修复后的文本和修复次数.
    """
    fixed = text
    count = 0

    style = _resolve_style(style)

    # 应用错误模式修复 (按体例开关选择规则组)
    for pattern, replacement in _patterns_for_style(style):
        new_text, n = pattern.subn(replacement, fixed)
        if n > 0:
            fixed = new_text
            count += n

    # 修复出版社格式: "××出版社××××年" → "××出版社××××年版"
    if _rule_enabled(style, 'publisher_year'):
        pub_fix = re.sub(
            r'([一-鿿]+出版社)\s*(\d{4})\s*年\s*$',
            r'\1\2年版',
            fixed
        )
        pub_fix = re.sub(
            r'([一-鿿]+出版社)\s*(\d{4})\s*年\s*[,，]',
            r'\1\2年版，',
            pub_fix
        )
        if pub_fix != fixed:
            count += 1
            fixed = pub_fix

    # 修复引注符号后的多余空格
    fixed = re.sub(r'^(\s*\d+)\s+(?=[一-鿿\[（])', r'\1 ', fixed)

    return fixed, count


def _get_footnotes_part(doc):
    """获取文档的脚注 Part 对象 (用于写回)."""
    FOOTNOTE_REL = ('http://schemas.openxmlformats.org/officeDocument/'
                    '2006/relationships/footnotes')
    for rel in doc.part.rels.values():
        if rel.reltype == FOOTNOTE_REL:
            return rel.target_part
    return None


def _write_footnote_part(fn_part, fn_xml) -> None:
    """将修改后的脚注 XML 写回 Part 对象.

    优先使用 Part._blob (python-docx 0.8.x ~ 1.x 一直稳定).
    若 _blob 不可用, 抛出 RuntimeError 并提示降级方案.

    Raises:
        RuntimeError: _blob 不可用且所有降级方案均失败.
    """
    try:
        from lxml import etree
        fn_part._blob = etree.tostring(
            fn_xml, xml_declaration=True, encoding='UTF-8', standalone=True)
    except AttributeError:
        import docx
        raise RuntimeError(
            f"脚注写入失败: python-docx Part._blob 不可用.\n"
            f"  python-docx 版本: {docx.__version__}\n"
            f"  降级方案: 文档保存后可运行 "
            f"python citation_formatter.py <文件> --fix-via-zip"
        ) from None


def write_footnotes_via_zipfile(docx_path: str) -> bool:
    """降级方案: 通过 ZIP 直操作修改已保存文档的 word/footnotes.xml.

    当 Part._blob 不可用时, 用此函数直接替换 ZIP 中的脚注 XML.
    先调用 format_all_footnotes 修复, 再重新读取修复后的 XML 写入 ZIP.

    Args:
        docx_path: 已保存的 .docx 文件路径.

    Returns:
        True 表示成功写入, False 表示文档中没有脚注.
    """
    import zipfile
    from lxml import etree
    from docx import Document

    # 1. 检查文档是否有脚注
    doc_check = Document(docx_path)
    if _get_footnotes_xml(doc_check) is None:
        return False

    # 2. 修复脚注 (format_all_footnotes 会修改内存中的 XML)
    #    然后从 Part._blob 重新解析得到修复后的 XML
    doc_fix = Document(docx_path)
    fn_xml_before = _get_footnotes_xml(doc_fix)
    if fn_xml_before is None:
        return False

    stats = format_all_footnotes(doc_fix, fix=True)
    if stats.get('write_error') or stats['fixed'] == 0:
        return False

    # 3. 从修复后的 Part._blob 重新解析 (避免拿到旧的 XML 引用)
    fn_xml_fixed = _get_footnotes_xml(doc_fix)
    if fn_xml_fixed is None:
        return False

    new_xml_bytes = etree.tostring(
        fn_xml_fixed, xml_declaration=True, encoding='UTF-8', standalone=True)

    # 4. ZIP 替换
    tmp_path = docx_path + '.tmp'
    with zipfile.ZipFile(docx_path, 'r') as zin:
        with zipfile.ZipFile(tmp_path, 'w', zipfile.ZIP_DEFLATED) as zout:
            for item in zin.infolist():
                if item.filename == 'word/footnotes.xml':
                    zout.writestr(item, new_xml_bytes)
                else:
                    zout.writestr(item, zin.read(item.filename))

    import os
    os.replace(tmp_path, docx_path)
    return True


def format_all_footnotes(doc: Document, fix: bool = True,
                         style=None) -> Dict:
    """格式化文档中所有脚注的引注内容.

    Args:
        doc: Document 对象.
        fix: True 时执行修复, False 仅检测.
        style: 注释体例 dict; None = 默认体例.
    """
    style = _resolve_style(style)
    fn_xml = _get_footnotes_xml(doc)
    if fn_xml is None:
        return {'total': 0, 'issues': 0, 'fixed': 0, 'checked': [],
                'message': '文档中没有脚注'}

    fn_part = _get_footnotes_part(doc)
    stats = {'total': 0, 'issues': 0, 'fixed': 0, 'checked': []}

    for fn_elem in fn_xml.findall(qn('w:footnote')):
        fn_id = int(fn_elem.get(qn('w:id'), '-1'))
        if fn_id <= 0:
            continue

        stats['total'] += 1
        fn_info = {
            'id': fn_id,
            'issues': [],
            'fixed_count': 0,
        }

        for p_elem in fn_elem.findall(qn('w:p')):
            run_elements = list(p_elem.findall(qn('w:r')))
            if not run_elements:
                continue

            # 找到 footnoteRef 位置 (自动编号所在 run 的索引)
            ref_run_idx = -1
            for idx, r_elem in enumerate(run_elements):
                if r_elem.find(qn('w:footnoteRef')) is not None:
                    ref_run_idx = idx
                    break

            # 分别收集脚注编号前后的文本
            prefix_text = ''    # footnoteRef 之前的文本 (如 "[")
            suffix_text = ''    # footnoteRef 之后的文本 (引注内容)

            for idx, r_elem in enumerate(run_elements):
                for t_elem in r_elem.findall(qn('w:t')):
                    if t_elem.text:
                        if idx < ref_run_idx:
                            prefix_text += t_elem.text
                        elif idx > ref_run_idx:
                            suffix_text += t_elem.text

            full_text = prefix_text + suffix_text
            if not suffix_text.strip():
                continue

            if fix:
                # 先修复, 再检测剩余问题. 修复不依赖检测门控 — auto_fix 的
                # 修复能力大于 check 的检测面时 (缺"年版"、终端句号等), 修复
                # 仍会执行, 检测结果反映修复后的剩余问题.
                fixed_text, fix_count = auto_fix_footnote(suffix_text, style=style)
                if fix_count > 0:
                    fn_info['fixed_count'] += fix_count
                    stats['fixed'] += fix_count

                    # 只修改 footnoteRef 之后的 run:
                    # 第一个 suffix run 写入修复后文本, 其余清空
                    # footnoteRef 及其之前的 run 原封不动
                    wrote_fixed = False
                    for idx, r_elem in enumerate(run_elements):
                        t_elems = r_elem.findall(qn('w:t'))
                        if idx > ref_run_idx and t_elems:
                            if not wrote_fixed:
                                t_elems[0].text = fixed_text
                                t_elems[0].set(qn('xml:space'), 'preserve')
                                wrote_fixed = True
                            else:
                                for t in t_elems:
                                    t.text = ''
                    suffix_text = fixed_text

            # 检测 (修复后文本; 未修复时为原始文本)
            issues = check_footnote(suffix_text)
            fn_info['issues'].extend(issues)
            stats['issues'] += len(issues)

        stats['checked'].append(fn_info)

    # Write modified XML back to part
    if fix and stats['fixed'] > 0 and fn_part is not None:
        try:
            _write_footnote_part(fn_part, fn_xml)
        except RuntimeError as e:
            stats.setdefault('write_error', str(e))
            print(f"  ⚠ {e}", file=sys.stderr)

    return stats


def main():
    parser = argparse.ArgumentParser(
        description='引注格式规范化工具 — 默认依据《法学引注手册》(2019), 支持自定义',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
示例:
  python citation_formatter.py 论文.docx --check          # 仅检查不修改
  python citation_formatter.py 论文.docx --fix             # 自动修复
  python citation_formatter.py 论文.docx -o 输出.docx      # 输出到新文件
  python citation_formatter.py --show-style                # 打印默认注释体例
  python citation_formatter.py --citation-rules 体例.json --show-style
        """,
    )
    parser.add_argument('input', nargs='?', help='输入 .docx 文件路径 '
                        '(使用 --show-style 时可不提供)')
    parser.add_argument('-o', '--output', help='输出文件路径')
    parser.add_argument('--check', action='store_true',
                        help='仅检查引注格式, 不修改')
    parser.add_argument('--fix', action='store_true',
                        help='自动修复检测到的格式问题')
    parser.add_argument('--citation-rules', default=None,
                        help='自定义注释体例 JSON 配置文件 '
                             '(见 references/citation-rules.md)')
    parser.add_argument('--preset', default=None,
                        help='预设注释体例: 法学引注手册 / 《法学家》《中外法学》注释体例 / 《中国法学》 / 《法商研究》 / 《法学研究》')
    parser.add_argument('--show-style', action='store_true',
                        help='打印当前注释体例并退出 (可与 --citation-rules/--preset 组合)')

    args = parser.parse_args()

    # —— 打印体例并退出 (不需要输入文件) ——
    if args.show_style:
        if not _HAS_CITATION_RULES:
            print('错误: 未找到 citation_rules 模块', file=sys.stderr)
            sys.exit(1)
        try:
            style = load_citation_rules(args.citation_rules, preset=args.preset)
        except (ValueError, FileNotFoundError) as e:
            print(f'错误: {e}', file=sys.stderr)
            sys.exit(1)
        print(describe_style(style))
        return

    if not args.input:
        print('错误: 请提供输入 .docx 文件路径', file=sys.stderr)
        sys.exit(1)

    input_path = Path(args.input)
    if not input_path.exists():
        print(f'错误: 文件不存在: {args.input}', file=sys.stderr)
        sys.exit(1)

    if not args.check and not args.fix:
        args.check = True  # 默认检查模式

    # —— 加载注释体例 (自定义/预设/默认) ——
    style = None
    if args.citation_rules or args.preset:
        if not _HAS_CITATION_RULES:
            print('错误: 未找到 citation_rules 模块 (需与 citation_formatter.py '
                  '同目录)', file=sys.stderr)
            sys.exit(1)
        try:
            style = load_citation_rules(args.citation_rules, preset=args.preset)
            src = args.citation_rules or f'预设 {args.preset}'
            print(f'已加载注释体例: {src}\n')
        except (ValueError, FileNotFoundError) as e:
            print(f'{{"error": "{e}", "error_type": "validation", '
                  f'"hint": "请检查注释体例配置文件格式 '
                  f'(见 references/citation-rules.md)"}}', file=sys.stderr)
            sys.exit(1)
    else:
        style = _resolve_style(None)

    doc = Document(str(input_path))

    # 先提取并显示所有脚注
    footnotes = extract_footnotes(doc)
    print(f'文档共有 {len(footnotes)} 条脚注\n')

    if args.check:
        print('=' * 60)
        print('引注格式检查报告')
        print('=' * 60)
        total_issues = 0

        for fn in footnotes:
            issues = check_footnote(fn['full_text'])
            total_issues += len(issues)
            if issues:
                print(f'\n[脚注 {fn["id"]}] {fn["full_text"][:80]}...')
                for iss in issues:
                    sev = {'high': '🔴', 'medium': '🟡', 'low': '🟢', 'info': 'ℹ️'}
                    print(f'  {sev.get(iss["severity"], "  ")} [{iss["type"]}] {iss["message"]}')
                    print(f'    建议: {iss["suggestion"]}')

        if total_issues == 0:
            print('\n✓ 未发现引注格式问题。')
        else:
            print(f'\n共发现 {total_issues} 个格式问题。')
            print('运行 --fix 可自动修复部分问题。')

    if args.fix:
        print('\n' + '=' * 60)
        print('自动修复引注格式')
        print('=' * 60)
        stats = format_all_footnotes(doc, fix=True, style=style)
        print(f'处理脚注: {stats["total"]} 条')
        print(f'发现问题: {stats["issues"]} 个')
        print(f'自动修复: {stats["fixed"]} 处')

        for fn_info in stats['checked']:
            if fn_info['issues']:
                print(f'\n[脚注 {fn_info["id"]}] '
                      f'{len(fn_info["issues"])} 个问题, '
                      f'{fn_info["fixed_count"]} 处已修复')

        output_path = args.output if args.output else str(input_path)
        doc.save(output_path)
        print(f'\n已保存至: {output_path}')

        unfixed = stats['issues'] - stats['fixed']
        if unfixed > 0:
            print(f'⚠ {unfixed} 个问题需手动处理 (复杂格式问题)')


if __name__ == '__main__':
    main()
