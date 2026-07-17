"""Integration tests: format documents and verify results."""

import pytest
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

from format_paper import format_document


def _get_run_font(run):
    """Extract font info from a docx run."""
    rPr = run._r.find(qn('w:rPr'))
    east = ascii_f = '?'
    if rPr is not None:
        rf = rPr.find(qn('w:rFonts'))
        if rf is not None:
            east = rf.get(qn('w:eastAsia'), '?')
            ascii_f = rf.get(qn('w:ascii'), '?')
    return {
        'eastAsia': east,
        'ascii': ascii_f,
        'size_pt': run.font.size / 12700 if run.font.size else 0,
        'bold': run.bold,
    }


def _get_paragraph_font(para):
    """Get font info from first run of a paragraph."""
    if not para.runs:
        return None
    return _get_run_font(para.runs[0])


def _get_alignment(para):
    """Get paragraph alignment as string."""
    return {
        WD_ALIGN_PARAGRAPH.CENTER: 'CENTER',
        WD_ALIGN_PARAGRAPH.LEFT: 'LEFT',
        WD_ALIGN_PARAGRAPH.RIGHT: 'RIGHT',
    }.get(para.alignment, 'OTHER')


def _has_indent(para):
    """Check if paragraph has first-line indent set."""
    pPr = para._p.find(qn('w:pPr'))
    if pPr is None:
        return False
    ind = pPr.find(qn('w:ind'))
    if ind is None:
        return False
    return ind.get(qn('w:firstLineChars')) is not None


@pytest.mark.integration
class TestTitleFormatting:
    """Verify title area formatting."""

    def test_title_font_and_size(self, minimal_docx_path, tmp_path):
        """Title should be 黑体, 14pt, no bold, centered."""
        out = str(tmp_path / 'out.docx')
        format_document(minimal_docx_path, out)

        doc = Document(out)
        # First non-empty paragraph = title
        for p in doc.paragraphs:
            if p.text.strip():
                info = _get_paragraph_font(p)
                assert info is not None
                assert info['eastAsia'] == '黑体', f'Expected 黑体, got {info["eastAsia"]}'
                assert info['size_pt'] == 14.0, f'Expected 14pt, got {info["size_pt"]}'
                assert info['bold'] is False
                assert _get_alignment(p) == 'CENTER'
                break


@pytest.mark.integration
class TestHeadingFormatting:
    """Verify heading level formatting."""

    def test_level1_heading(self, minimal_docx_path, tmp_path):
        """Level 1 heading: 宋体, 12pt, bold, centered."""
        out = str(tmp_path / 'out.docx')
        format_document(minimal_docx_path, out)

        doc = Document(out)
        # Find paragraph with 一、
        for p in doc.paragraphs:
            if p.text.strip().startswith('一、'):
                info = _get_paragraph_font(p)
                assert info['eastAsia'] == '宋体'
                assert info['size_pt'] == 12.0
                assert info['bold'] is True
                assert _get_alignment(p) == 'CENTER'
                return
        pytest.fail('No level-1 heading found')

    def test_level2_heading(self, tmp_path):
        """Level 2 heading: 楷体, 12pt, bold, indent."""
        doc = Document()
        doc.add_paragraph('测试题目')
        doc.add_paragraph('一、概述')
        doc.add_paragraph('（一）研究背景')
        doc.add_paragraph('正文内容。')
        inp = str(tmp_path / 'in.docx')
        doc.save(inp)

        out = str(tmp_path / 'out.docx')
        format_document(inp, out)

        result = Document(out)
        for p in result.paragraphs:
            if p.text.strip().startswith('（一）'):
                info = _get_paragraph_font(p)
                assert info['eastAsia'] == '楷体'
                assert info['size_pt'] == 12.0
                assert info['bold'] is True
                assert _has_indent(p)
                return
        pytest.fail('No level-2 heading found')

    def test_level3_heading(self, tmp_path):
        """Level 3 heading: 宋体, 10.5pt, bold, indent."""
        doc = Document()
        doc.add_paragraph('测试题目')
        doc.add_paragraph('一、概述')
        doc.add_paragraph('1. 研究方法')
        doc.add_paragraph('正文内容。')
        inp = str(tmp_path / 'in.docx')
        doc.save(inp)

        out = str(tmp_path / 'out.docx')
        format_document(inp, out)

        result = Document(out)
        for p in result.paragraphs:
            if p.text.strip().startswith('1. '):
                info = _get_paragraph_font(p)
                assert info['eastAsia'] == '宋体'
                assert info['size_pt'] == 10.5
                assert info['bold'] is True
                return
        pytest.fail('No level-3 heading found')


@pytest.mark.integration
class TestBodyFormatting:
    """Verify body text formatting."""

    def test_body_font(self, minimal_docx_path, tmp_path):
        """Body text: 宋体, 10.5pt, no bold."""
        out = str(tmp_path / 'out.docx')
        format_document(minimal_docx_path, out)

        doc = Document(out)
        for p in doc.paragraphs:
            if '正文内容' in p.text:
                info = _get_paragraph_font(p)
                assert info['eastAsia'] == '宋体'
                assert info['size_pt'] == 10.5
                assert info['bold'] is False
                return
        pytest.fail('No body paragraph found')

    def test_ascii_font_is_song(self, minimal_docx_path, tmp_path):
        """Arabic digits should use 宋体, not Times New Roman."""
        out = str(tmp_path / 'out.docx')
        format_document(minimal_docx_path, out)

        doc = Document(out)
        for p in doc.paragraphs:
            if '正文内容' in p.text:
                info = _get_paragraph_font(p)
                assert info['ascii'] == '宋体', \
                    f'Arabic numerals should be 宋体, got {info["ascii"]}'
                return
        pytest.fail('No body paragraph found')

    def test_body_indent_applied(self, minimal_docx_path, tmp_path):
        """--body-indent 2 should set first-line indent."""
        out = str(tmp_path / 'out.docx')
        format_document(minimal_docx_path, out, body_indent=2)

        doc = Document(out)
        for p in doc.paragraphs:
            if '正文内容' in p.text:
                assert _has_indent(p), 'Body should have indent with --body-indent 2'
                return

    def test_body_indent_not_applied(self, minimal_docx_path, tmp_path):
        """Default (no --body-indent) should NOT indent body."""
        out = str(tmp_path / 'out.docx')
        format_document(minimal_docx_path, out, body_indent=0)

        doc = Document(out)
        for p in doc.paragraphs:
            if '正文内容' in p.text:
                assert not _has_indent(p), \
                    'Body should NOT have indent without --body-indent'
                return


@pytest.mark.integration
class TestAbstractFormatting:
    """Verify abstract and keywords formatting."""

    def test_abstract_kaiti(self, docx_with_abstract, tmp_path):
        """Abstract should be 楷体, 12pt, no bold."""
        out = str(tmp_path / 'out.docx')
        format_document(docx_with_abstract, out)

        doc = Document(out)
        for p in doc.paragraphs:
            if '摘要' in p.text:
                info = _get_paragraph_font(p)
                assert info['eastAsia'] == '楷体', \
                    f'Expected 楷体, got {info["eastAsia"]}'
                assert info['size_pt'] == 12.0
                assert info['bold'] is False
                return
        pytest.fail('No abstract paragraph found')

    def test_keywords_kaiti(self, docx_with_abstract, tmp_path):
        """Keywords should be 楷体, 12pt, no bold."""
        out = str(tmp_path / 'out.docx')
        format_document(docx_with_abstract, out)

        doc = Document(out)
        for p in doc.paragraphs:
            if '关键词' in p.text:
                info = _get_paragraph_font(p)
                assert info['eastAsia'] == '楷体', \
                    f'Expected 楷体, got {info["eastAsia"]}'
                assert info['size_pt'] == 12.0
                assert info['bold'] is False
                return
        pytest.fail('No keywords paragraph found')

    def test_abstract_no_indent(self, docx_with_abstract, tmp_path):
        """Abstract should NOT have first-line indent."""
        out = str(tmp_path / 'out.docx')
        format_document(docx_with_abstract, out)

        doc = Document(out)
        for p in doc.paragraphs:
            if '摘要' in p.text:
                assert not _has_indent(p), 'Abstract should not be indented'
                return


@pytest.mark.integration
class TestTitleRange:
    """Verify title area is limited to 2 paragraphs."""

    def test_title_range_max_two(self, tmp_path):
        """Only first 2 non-empty paragraphs → title area. Rest → body."""
        doc = Document()
        doc.add_paragraph('笔记1：适用规则之争')       # → title
        doc.add_paragraph('笔记2：损失界定')            # → subtitle (title area)
        doc.add_paragraph('笔记3：违约金与解除权')      # → body (beyond limit)
        doc.add_paragraph('')                           # blank
        doc.add_paragraph('竞业限制合同的违约金调整规则研究')  # → body
        doc.add_paragraph('')
        doc.add_paragraph('一、引言')
        doc.add_paragraph('正文内容。')
        inp = str(tmp_path / 'in.docx')
        doc.save(inp)

        out = str(tmp_path / 'out.docx')
        format_document(inp, out)

        result = Document(out)
        paras = [p for p in result.paragraphs if p.text.strip()]

        # Paragraph 0 (笔记1) → title (黑体)
        info0 = _get_paragraph_font(paras[0])
        assert info0['eastAsia'] == '黑体'
        assert info0['size_pt'] == 14.0

        # Paragraph 1 (笔记2) → also title area (within 2-para limit)
        info1 = _get_paragraph_font(paras[1])
        assert info1['eastAsia'] == '黑体'

        # Paragraph 2 (笔记3) → should be BODY text, NOT title
        info2 = _get_paragraph_font(paras[2])
        assert info2['eastAsia'] == '宋体', \
            f'Paragraph 3 should be body (宋体), got {info2["eastAsia"]}'
        assert info2['size_pt'] == 10.5, \
            f'Paragraph 3 should be body (10.5pt), got {info2["size_pt"]}'
