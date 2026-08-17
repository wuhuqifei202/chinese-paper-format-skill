"""Shared pytest fixtures for chinese-paper-format-skill tests."""

import io
import zipfile
from pathlib import Path

import pytest
from docx import Document
from docx.shared import Pt
from lxml import etree

# OOXML namespaces
WML = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
REL = 'http://schemas.openxmlformats.org/package/2006/relationships'
FOOTNOTE_REL = 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/footnotes'
CT = 'http://schemas.openxmlformats.org/package/2006/content-types'


def _add_footnote_reference(para_element, footnote_id: int):
    """Add a footnote reference run to a paragraph XML element."""
    r_elem = etree.SubElement(para_element, f'{{{WML}}}r')
    rPr = etree.SubElement(r_elem, f'{{{WML}}}rPr')
    va = etree.SubElement(rPr, f'{{{WML}}}vertAlign')
    va.set(f'{{{WML}}}val', 'superscript')
    fn_ref = etree.SubElement(r_elem, f'{{{WML}}}footnoteReference')
    fn_ref.set(f'{{{WML}}}id', str(footnote_id))


def _build_footnote(fn_id: int, text: str):
    """Build a footnote XML element with given id and text."""
    fn = etree.Element(f'{{{WML}}}footnote')
    fn.set(f'{{{WML}}}id', str(fn_id))
    p = etree.SubElement(fn, f'{{{WML}}}p')

    # Left bracket
    r0 = etree.SubElement(p, f'{{{WML}}}r')
    t0 = etree.SubElement(r0, f'{{{WML}}}t')
    t0.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    t0.text = '['

    # Auto-number
    r1 = etree.SubElement(p, f'{{{WML}}}r')
    etree.SubElement(r1, f'{{{WML}}}footnoteRef')

    # Right bracket + citation text
    r2 = etree.SubElement(p, f'{{{WML}}}r')
    t2 = etree.SubElement(r2, f'{{{WML}}}t')
    t2.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    t2.text = ']' + text

    return fn


def _inject_footnotes(docx_bytes: bytes, footnote_xml_elements: list) -> bytes:
    """Inject footnote XML elements into a .docx ZIP archive."""
    with zipfile.ZipFile(io.BytesIO(docx_bytes), 'r') as zin:
        out_buf = io.BytesIO()
        with zipfile.ZipFile(out_buf, 'w', zipfile.ZIP_DEFLATED) as zout:
            for item in zin.infolist():
                data = zin.read(item.filename)

                if item.filename == 'word/document.xml':
                    doc_xml = etree.fromstring(data)
                    body = doc_xml.find(f'{{{WML}}}body')
                    first_p = body.find(f'{{{WML}}}p')
                    if first_p is not None:
                        for fid in range(1, len(footnote_xml_elements) + 1):
                            _add_footnote_reference(first_p, fid)
                    data = etree.tostring(doc_xml, xml_declaration=True,
                                          encoding='UTF-8', standalone=True)

                elif item.filename == 'word/_rels/document.xml.rels':
                    rels_xml = etree.fromstring(data)
                    existing = [r for r in rels_xml
                                if r.get('Type') == FOOTNOTE_REL]
                    if not existing:  # 查重: 重复 rel → Word 判损坏
                        fn_rel = etree.SubElement(rels_xml, f'{{{REL}}}Relationship')
                        fn_rel.set('Id', 'rIdFootnotes')
                        fn_rel.set('Type', FOOTNOTE_REL)
                        fn_rel.set('Target', 'footnotes.xml')
                    data = etree.tostring(rels_xml, xml_declaration=True,
                                          encoding='UTF-8', standalone=True)

                elif item.filename in ('[Content_Types].xml', 'word/footnotes.xml'):
                    continue  # 末尾统一写入, 避免 Duplicate name 警告

                zout.writestr(item, data)

            # Build footnotes.xml
            fns_root = etree.Element(f'{{{WML}}}footnotes',
                                      nsmap={'w': WML})
            # Separator footnotes (必须带 w:type, 否则 Word 按普通脚注解析 id=-1 → 损坏)
            for sep_id, sep_tag in [('-1', 'separator'), ('0', 'continuationSeparator')]:
                sep = etree.SubElement(fns_root, f'{{{WML}}}footnote')
                sep.set(f'{{{WML}}}type', sep_tag)
                sep.set(f'{{{WML}}}id', sep_id)
                sp = etree.SubElement(sep, f'{{{WML}}}p')
                sr = etree.SubElement(sp, f'{{{WML}}}r')
                etree.SubElement(sr, f'{{{WML}}}{sep_tag}')

            for elem in footnote_xml_elements:
                fns_root.append(elem)

            zout.writestr('word/footnotes.xml',
                          etree.tostring(fns_root, xml_declaration=True,
                                         encoding='UTF-8', standalone=True))

            # Update Content_Types
            ct_xml = etree.fromstring(zin.read('[Content_Types].xml'))
            override = etree.SubElement(ct_xml, f'{{{CT}}}Override')
            override.set('PartName', '/word/footnotes.xml')
            override.set('ContentType',
                         'application/vnd.openxmlformats-officedocument.wordprocessingml.footnotes+xml')
            zout.writestr('[Content_Types].xml',
                          etree.tostring(ct_xml, xml_declaration=True,
                                         encoding='UTF-8', standalone=True))

    return out_buf.getvalue()


# ── Fixtures ──────────────────────────────────────────────────────────


@pytest.fixture
def minimal_docx_path(tmp_path):
    """Create a minimal .docx: title + level-1 heading + body text."""
    doc = Document()
    doc.add_paragraph('测试论文题目')
    doc.add_paragraph('一、引言')
    doc.add_paragraph('这是正文内容，包含一些中文文本和数字123。')

    path = str(tmp_path / 'minimal.docx')
    doc.save(path)
    return path


@pytest.fixture
def docx_with_abstract(tmp_path):
    """Create a .docx with abstract and keywords."""
    doc = Document()
    doc.add_paragraph('合同解除权行使规则研究')
    doc.add_paragraph('【摘要】本文研究了合同解除权的行使要件与法律效果。')
    doc.add_paragraph('【关键词】合同解除权；解除条件；法律效果')
    doc.add_paragraph('一、绪论')
    doc.add_paragraph('这是正文内容。')

    path = str(tmp_path / 'with_abstract.docx')
    doc.save(path)
    return path


@pytest.fixture
def docx_with_footnotes(tmp_path):
    """Create a .docx with 3 footnotes containing citation errors."""
    doc = Document()
    doc.add_paragraph('试论人工智能的侵权责任')
    doc.add_paragraph('一、问题的提出')
    p1 = doc.add_paragraph('人工智能侵权责任问题日益突出。')
    p1.runs[0].font.size = Pt(10.5)

    # Save and inject footnotes
    buf = io.BytesIO()
    doc.save(buf)
    raw = buf.getvalue()

    fn1 = _build_footnote(1, '示例学者， 《示例法学研究》, 示例出版社 2020 年, 第 125 页.')
    fn2 = _build_footnote(2, '示例学者：《示例法学通论》，示例出版社 2019 年版，p.89.')
    fn3 = _build_footnote(3, '《民法典》第68条。')

    modified = _inject_footnotes(raw, [fn1, fn2, fn3])

    path = str(tmp_path / 'with_footnotes.docx')
    with open(path, 'wb') as f:
        f.write(modified)
    return path


@pytest.fixture
def docx_with_legal_citations(tmp_path):
    """Create a .docx with legal document and case citations."""
    doc = Document()
    doc.add_paragraph('法律文件与案例引用测试')
    doc.add_paragraph('一、法律文件引用')
    doc.add_paragraph('测试规范性文件引用。')

    buf = io.BytesIO()
    doc.save(buf)
    raw = buf.getvalue()

    fn1 = _build_footnote(1, '《示例规范性文件》，示例发[2020]19号。')
    fn2 = _build_footnote(2, '示例案例，示例市人民法院[2021]示例初字第142号民事判决书。')
    fn3 = _build_footnote(3, '《示例司法解释》，示例发[2018]1号，第100条。')

    modified = _inject_footnotes(raw, [fn1, fn2, fn3])

    path = str(tmp_path / 'with_legal.docx')
    with open(path, 'wb') as f:
        f.write(modified)
    return path


@pytest.fixture
def docx_with_author_line(tmp_path):
    """题目 + 作者行 + 摘要/关键词 (复现 BUG-008: 作者行后摘要被误判为正文)."""
    doc = Document()
    doc.add_paragraph('算法歧视的法律规制研究')
    doc.add_paragraph('人工智能与法律的交叉研究课题组')
    doc.add_paragraph('【摘要】算法歧视是人工智能治理面临的核心问题之一。')
    doc.add_paragraph('【关键词】算法歧视；算法治理')
    doc.add_paragraph('一、问题的提出')
    doc.add_paragraph('算法在信贷、招聘等领域的应用带来了歧视风险。')

    path = str(tmp_path / 'with_author_line.docx')
    doc.save(path)
    return path


@pytest.fixture
def docx_with_missing_year(tmp_path):
    """脚注出版社缺"年版" (复现 T05 fn4)."""
    doc = Document()
    doc.add_paragraph('论网络服务提供者的注意义务')
    doc.add_paragraph('一、注意义务的判断标准')
    doc.add_paragraph('网络服务提供者的注意义务应当综合判断。')

    buf = io.BytesIO()
    doc.save(buf)
    raw = buf.getvalue()

    fn1 = _build_footnote(1, '示例学者：《示例法学总论》，示例出版社 2017 年')

    modified = _inject_footnotes(raw, [fn1])

    path = str(tmp_path / 'with_missing_year.docx')
    with open(path, 'wb') as f:
        f.write(modified)
    return path


@pytest.fixture
def docx_with_trailing_period(tmp_path):
    """脚注终端英文句号 (复现 T10 fn1)."""
    doc = Document()
    doc.add_paragraph('算法歧视的法律规制研究')
    doc.add_paragraph('一、问题的提出')
    doc.add_paragraph('算法在信贷领域的应用带来了歧视风险。')

    buf = io.BytesIO()
    doc.save(buf)
    raw = buf.getvalue()

    fn1 = _build_footnote(1, '参见示例学者：《示例算法法学》，载《示例法学研究》2019年第4期，第50页.')

    modified = _inject_footnotes(raw, [fn1])

    path = str(tmp_path / 'with_trailing_period.docx')
    with open(path, 'wb') as f:
        f.write(modified)
    return path
