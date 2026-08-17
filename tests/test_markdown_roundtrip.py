"""markdown 转换层测试: docx → md → docx 往返 (格式转换层).

覆盖: md 结构约定 / 脚注定义 / 排版重建 / 作者行与摘要分类 /
md 层引注修复 / 连续脚注定义行解析.
"""

import pytest
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

from docx2md import docx_to_markdown
from md2docx import parse_markdown, markdown_to_docx
from run_pipeline import fix_citations_in_markdown


def _get_run_font(run):
    """Extract font info from a docx run."""
    rPr = run._r.find(qn('w:rPr'))
    east = '?'
    if rPr is not None:
        rf = rPr.find(qn('w:rFonts'))
        if rf is not None:
            east = rf.get(qn('w:eastAsia'), '?')
    return {
        'eastAsia': east,
        'size_pt': run.font.size / 12700 if run.font.size else 0,
        'bold': run.bold,
    }


def _has_indent(para):
    """Check if paragraph has first-line indent set."""
    pPr = para._p.find(qn('w:pPr'))
    if pPr is None:
        return False
    ind = pPr.find(qn('w:ind'))
    if ind is None:
        return False
    return ind.get(qn('w:firstLineChars')) is not None


def _is_centered(para):
    return para.alignment == WD_ALIGN_PARAGRAPH.CENTER


def _run_roundtrip(docx_path, tmp_path, **kwargs):
    """docx → md → docx, 返回 (md_text, rebuilt_path)."""
    md_text = docx_to_markdown(docx_path)
    out = str(tmp_path / 'rebuilt.docx')
    stats = markdown_to_docx(md_text, out, **kwargs)
    return md_text, out, stats


@pytest.mark.integration
class TestDocxToMd:
    """docx → markdown 转换."""

    def test_title_prefix(self, minimal_docx_path):
        """题目应以 '# ' 开头."""
        md = docx_to_markdown(minimal_docx_path)
        assert md.startswith('# 测试论文题目'), f'题目前缀错误: {md[:30]!r}'

    def test_heading_prefixes(self, docx_with_author_line):
        """一级/二级标题前缀: ## / ###."""
        md = docx_to_markdown(docx_with_author_line)
        assert '\n## 一、问题的提出' in md
        assert '# 算法歧视的法律规制研究' in md

    def test_author_and_abstract_preserved(self, docx_with_author_line):
        """作者行/摘要/关键词原样保留 (无 md 标记)."""
        md = docx_to_markdown(docx_with_author_line)
        assert '人工智能与法律的交叉研究课题组' in md
        assert '【摘要】算法歧视是人工智能治理面临的核心问题之一' in md
        assert '【关键词】算法歧视' in md

    def test_footnote_defs(self, docx_with_footnotes):
        """脚注以 [^n]: 形式置于文末."""
        md = docx_to_markdown(docx_with_footnotes)
        assert '[^1]: ' in md and '[^2]: ' in md and '[^3]: ' in md
        # 脚注编号残留 "[]" 应被剥离
        assert '[]示例学者' not in md
        assert '示例学者' in md

    def test_footnote_ref_marker(self, docx_with_footnotes):
        """正文段落末尾应有 [^n] 引用标记."""
        md = docx_to_markdown(docx_with_footnotes)
        assert '[^1]' in md
        # 引用标记出现在脚注定义之前
        assert md.index('[^1]') < md.index('[^1]: ')


@pytest.mark.integration
class TestMdToDocx:
    """markdown → docx 重建."""

    def test_roundtrip_formats(self, minimal_docx_path, tmp_path):
        """往返后排版规则应正确套用."""
        md, out, stats = _run_roundtrip(minimal_docx_path, tmp_path)
        assert stats['title'] == 1
        assert stats['headings_l1'] >= 1
        assert stats['body'] >= 1

        doc = Document(out)
        paras = [p for p in doc.paragraphs if p.text.strip()]
        # 题目: 黑体 14pt 居中
        info = _get_run_font(paras[0].runs[0])
        assert info['eastAsia'] == '黑体'
        assert info['size_pt'] == 14.0
        assert _is_centered(paras[0])
        # 一级: 宋体 12pt 加粗
        for p in paras:
            if p.text.strip().startswith('一、'):
                info = _get_run_font(p.runs[0])
                assert info['eastAsia'] == '宋体'
                assert info['size_pt'] == 12.0
                assert info['bold'] is True
                break
        # 正文: 宋体 10.5pt 缩进
        for p in paras:
            if '正文内容' in p.text:
                info = _get_run_font(p.runs[0])
                assert info['eastAsia'] == '宋体'
                assert info['size_pt'] == 10.5
                assert _has_indent(p)
                break

    def test_roundtrip_footnotes(self, docx_with_footnotes, tmp_path):
        """往返后脚注应保留且可读."""
        md, out, stats = _run_roundtrip(docx_with_footnotes, tmp_path)
        assert stats['footnotes'] == 3

        doc = Document(out)
        from citation_formatter import extract_footnotes
        fns = extract_footnotes(doc)
        assert len(fns) == 3
        texts = {fn['id']: fn['full_text'] for fn in fns}
        assert '示例法学研究' in texts[1]

    def test_roundtrip_author_abstract(self, docx_with_author_line, tmp_path):
        """往返后作者行/摘要/关键词分类正确."""
        md, out, stats = _run_roundtrip(docx_with_author_line, tmp_path)
        assert stats['author'] == 1
        assert stats['abstract'] == 1
        assert stats['keywords'] == 1

        doc = Document(out)
        for p in doc.paragraphs:
            t = p.text.strip()
            if '课题组' in t:
                info = _get_run_font(p.runs[0])
                assert info['eastAsia'] == '黑体'
                assert info['size_pt'] == 14.0
                assert _is_centered(p)
            elif '【摘要】' in t:
                info = _get_run_font(p.runs[0])
                assert info['eastAsia'] == '楷体'
                assert info['size_pt'] == 12.0
                assert not _has_indent(p)

    def test_parse_markdown_consecutive_fndefs(self):
        """连续脚注定义行 (无空行分隔) 不应互相吞并."""
        md = ('# 测试\n\n正文[^1]\n\n[^1]: 第一条脚注\n'
              '[^2]: 第二条脚注\n[^3]: 第三条脚注\n')
        blocks, footnotes = parse_markdown(md)
        assert len(footnotes) == 3
        assert footnotes[1] == '第一条脚注'
        assert footnotes[2] == '第二条脚注'
        assert footnotes[3] == '第三条脚注'

    def test_parse_markdown_missing_title(self, tmp_path):
        """缺少 '# ' 题目应报错."""
        with pytest.raises(ValueError):
            markdown_to_docx('一、标题\n\n正文\n', str(tmp_path / 'out.docx'))

    def test_rebuild_ignores_unknown_refs(self, tmp_path):
        """引用了但未定义的脚注应补齐空定义 (避免 Word 报错)."""
        md = '# 测试\n\n正文[^9]\n'
        out = str(tmp_path / 'out.docx')
        stats = markdown_to_docx(md, out)
        assert stats['footnotes'] == 1


@pytest.mark.pure
class TestMdLayerCitationFix:
    """md 层引注修复 (run_pipeline.fix_citations_in_markdown)."""

    def test_fix_year_and_period(self):
        """缺年版与终端英文句号应修复."""
        md = ('# 测试\n\n[^1]: 示例学者：《示例法学总论》，示例出版社 2017 年\n'
              '[^2]: 参见示例学者：《示例算法法学》，第50页.\n')
        fixed, count = fix_citations_in_markdown(md)
        assert count >= 2
        assert '示例出版社2017年版' in fixed
        assert '第50页。' in fixed
        # 修复后仍是合法脚注定义行
        assert '\n[^1]: 示例学者：《示例法学总论》，示例出版社2017年版' in fixed

    def test_roundtrip_fix_applies(self, docx_with_missing_year, tmp_path):
        """docx → md → (修复) → docx, 缺年版应出现在重建脚注中."""
        md = docx_to_markdown(docx_with_missing_year)
        md_fixed, count = fix_citations_in_markdown(md)
        assert count >= 1
        out = str(tmp_path / 'out.docx')
        markdown_to_docx(md_fixed, out)

        doc = Document(out)
        from citation_formatter import extract_footnotes
        fns = extract_footnotes(doc)
        assert '示例出版社2017年版' in fns[0]['full_text']
