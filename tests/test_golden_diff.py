"""Golden file regression tests — 格式化结果不能意外改变.

每次修改代码后，重新格式化相同的输入论文，逐段与黄金输出对比。
任何意外的格式变化都会导致测试失败。
"""

import sys
from pathlib import Path

import pytest
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

from format_paper import format_document
from citation_formatter import extract_footnotes

SKILL_DIR = Path(__file__).resolve().parent.parent
GOLDEN_DIR = SKILL_DIR / 'evals' / 'golden'

# ── 每个 golden case 的定义 ──────────────────────────────────────────────
GOLDEN_CASES = [
    {
        'id': '担保型以物抵债',
        'input': GOLDEN_DIR / '担保型以物抵债' / 'input.docx',
        'expected': GOLDEN_DIR / '担保型以物抵债' / 'expected.docx',
        'body_indent': 2,
        'fix_citations': True,
    },
    {
        'id': '竞业限制',
        'input': GOLDEN_DIR / '竞业限制' / 'input.docx',
        'expected': GOLDEN_DIR / '竞业限制' / 'expected.docx',
        'body_indent': 2,
        'fix_citations': False,
    },
]


def _paragraph_signature(para) -> dict:
    """提取段落格式签名: 仅取第一个 run 的字体属性 + 对齐."""
    if not para.runs:
        return {'text': '', 'font': None, 'align': None}
    r = para.runs[0]
    rPr = r._r.find(qn('w:rPr'))
    ea = ascii_f = '?'
    if rPr is not None:
        rf = rPr.find(qn('w:rFonts'))
        if rf is not None:
            ea = rf.get(qn('w:eastAsia'), '?')
            ascii_f = rf.get(qn('w:ascii'), '?')
    return {
        'text': para.text.strip()[:60],
        'font': {
            'eastAsia': ea,
            'ascii': ascii_f,
            'size_pt': round(r.font.size / 12700, 1) if r.font.size else 0,
            'bold': r.bold or False,
        },
        'align': {
            WD_ALIGN_PARAGRAPH.CENTER: 'CENTER',
            WD_ALIGN_PARAGRAPH.LEFT: 'LEFT',
            WD_ALIGN_PARAGRAPH.RIGHT: 'RIGHT',
        }.get(para.alignment, 'OTHER'),
    }


def _paragraphs_diff(actual_doc, expected_doc) -> list[str]:
    """逐段对比两个文档，返回差异列表."""
    diffs = []
    a_paras = [p for p in actual_doc.paragraphs if p.text.strip()]
    e_paras = [p for p in expected_doc.paragraphs if p.text.strip()]

    if len(a_paras) != len(e_paras):
        diffs.append(
            f'段落数不一致: actual={len(a_paras)}, expected={len(e_paras)}'
        )
        # Continue comparing what we can
        max_len = min(len(a_paras), len(e_paras))
    else:
        max_len = len(a_paras)

    for i in range(max_len):
        a_sig = _paragraph_signature(a_paras[i])
        e_sig = _paragraph_signature(e_paras[i])
        if a_sig != e_sig:
            diffs.append(
                f'[{i}] 段落格式变化:\n'
                f'  actual:   {a_sig}\n'
                f'  expected: {e_sig}'
            )
    return diffs


def _footnotes_diff(actual_doc, expected_doc) -> list[str]:
    """对比脚注文本，返回差异列表."""
    diffs = []
    a_fns = extract_footnotes(actual_doc)
    e_fns = extract_footnotes(expected_doc)

    if len(a_fns) != len(e_fns):
        diffs.append(
            f'脚注数不一致: actual={len(a_fns)}, expected={len(e_fns)}'
        )

    # Compare footnote text by ID
    a_by_id = {f['id']: f['full_text'] for f in a_fns}
    e_by_id = {f['id']: f['full_text'] for f in e_fns}

    all_ids = set(a_by_id.keys()) | set(e_by_id.keys())
    for fid in sorted(all_ids):
        a_text = a_by_id.get(fid, '(missing)')
        e_text = e_by_id.get(fid, '(missing)')
        if a_text != e_text:
            # Only show first 80 chars
            diffs.append(
                f'[Fn {fid}] 脚注内容变化:\n'
                f'  actual:   {a_text[:80]}\n'
                f'  expected: {e_text[:80]}'
            )
    return diffs


@pytest.mark.integration
@pytest.mark.parametrize('case', GOLDEN_CASES, ids=[c['id'] for c in GOLDEN_CASES])
class TestGoldenDiff:
    """Golden file diff — 格式化结果必须与黄金输出一致."""

    def test_paragraph_formatting(self, case, tmp_path):
        """段落格式不能意外变化."""
        actual_output = str(tmp_path / 'actual.docx')
        format_document(
            str(case['input']), actual_output,
            body_indent=case['body_indent']
        )

        actual_doc = Document(actual_output)
        expected_doc = Document(str(case['expected']))

        diffs = _paragraphs_diff(actual_doc, expected_doc)
        if diffs:
            pytest.fail(
                f'\n{"=" * 60}\n'
                f'Golden diff 失败: {case["id"]}\n'
                f'共 {len(diffs)} 处差异:\n'
                + '\n'.join(diffs[:20])  # cap at 20 diffs
            )

    def test_footnote_text(self, case, tmp_path):
        """脚注文本不能意外变化."""
        if not case['fix_citations']:
            pytest.skip('该 golden case 不需要引注修复')

        actual_output = str(tmp_path / 'actual.docx')
        format_document(
            str(case['input']), actual_output,
            body_indent=case['body_indent']
        )
        # Apply citation fix to match expected
        from citation_formatter import format_all_footnotes
        actual_doc = Document(actual_output)
        format_all_footnotes(actual_doc, fix=True)
        actual_doc.save(actual_output)

        actual_doc2 = Document(actual_output)
        expected_doc = Document(str(case['expected']))

        diffs = _footnotes_diff(actual_doc2, expected_doc)
        if diffs:
            pytest.fail(
                f'\n{"=" * 60}\n'
                f'Footnote diff 失败: {case["id"]}\n'
                + '\n'.join(diffs[:15])
            )
