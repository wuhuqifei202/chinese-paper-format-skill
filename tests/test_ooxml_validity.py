"""OOXML 结构合法性回归测试 — Word 拒开文档的三个已确认损坏点.

历史: 格式化输出被 Word 拒绝打开 (普通模式), OpenAndRepair 修复可恢复.
根因三处:
1. w:ind 被插入到 pPr 的 w:rPr 之后 (CT_PPr 元素顺序违规)
2. footnotes.xml 的 separator(-1)/continuationSeparator(0) 缺少 w:type 属性,
   无 type 时 Word 按普通脚注解析, id=-1 非法 → 拒开
3. document.xml.rels 中重复 footnotes 关系 (rId5 + rIdFootnotes 同 target)

任何一个复现 → Word 普通模式拒开文档. 本文件锁定这三个点.
"""

import io
import re
import zipfile

import pytest
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from lxml import etree

from format_paper import (_insert_pPr_ordered, set_first_line_indent,
                          set_line_spacing, format_document)
from md2docx import _build_footnotes_xml

WML = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
W = '{%s}' % WML
FOOTNOTE_REL = ('http://schemas.openxmlformats.org/officeDocument/'
                '2006/relationships/footnotes')

# CT_PPr 全序 (ECMA-376), 用于顺序判定
_CT_PPR = ('pStyle', 'keepNext', 'keepLines', 'pageBreakBefore', 'framePr',
           'widowControl', 'numPr', 'suppressLineNumbers', 'pBdr', 'shd',
           'tabs', 'suppressAutoHyphens', 'kinsoku', 'wordWrap',
           'overflowPunct', 'topLinePunct', 'autoSpaceDE', 'autoSpaceDN',
           'bidi', 'adjustRightInd', 'snapToGrid', 'spacing', 'ind',
           'contextualSpacing', 'mirrorIndents', 'suppressOverlap', 'jc',
           'textDirection', 'textAlignment', 'textboxTightWrap', 'outlineLvl',
           'divId', 'cnfStyle', 'rPr', 'sectPr', 'pPrChange')


def _pPr_with_children(tags: list):
    """构造含指定子元素的 pPr (空元素即可, 仅测顺序)."""
    pPr = etree.Element(W + 'pPr')
    for tag in tags:
        etree.SubElement(pPr, W + tag)
    return pPr


@pytest.mark.pure
class TestInsertIndOrder:
    """w:ind 插入位置必须符合 CT_PPr 顺序 (在 rPr 之前)."""

    def test_empty_pPr(self):
        pPr = _pPr_with_children([])
        ind = etree.Element(W + 'ind')
        _insert_pPr_ordered(pPr, ind)
        assert pPr[0] is ind

    def test_ind_after_spacing_only(self):
        pPr = _pPr_with_children(['spacing'])
        ind = etree.Element(W + 'ind')
        _insert_pPr_ordered(pPr, ind)
        assert [etree.QName(c).localname for c in pPr] == ['spacing', 'ind']

    def test_ind_before_jc(self):
        """pPr 含 jc (无 rPr): ind 插在 jc 前."""
        pPr = _pPr_with_children(['spacing', 'jc'])
        ind = etree.Element(W + 'ind')
        _insert_pPr_ordered(pPr, ind)
        assert [etree.QName(c).localname for c in pPr] == \
            ['spacing', 'ind', 'jc']

    def test_ind_before_rPr(self):
        """pPr 同时含 jc 与 rPr (最易踩坑形态): ind 必须插在 rPr 前.

        python-docx 的 get_or_add_ind() 会把 ind 插到 rPr 之后 → Word 拒开.
        """
        pPr = _pPr_with_children(['spacing', 'jc', 'rPr'])
        ind = etree.Element(W + 'ind')
        _insert_pPr_ordered(pPr, ind)
        names = [etree.QName(c).localname for c in pPr]
        assert names == ['spacing', 'ind', 'jc', 'rPr']

    def test_sequence_violation_check(self):
        """产物整体必须无 CT_PPr 顺序违规."""
        pPr = _pPr_with_children(['spacing', 'jc', 'rPr'])
        ind = etree.Element(W + 'ind')
        _insert_pPr_ordered(pPr, ind)
        pos = {name: i for i, name in enumerate(_CT_PPR)}
        last = -1
        for c in pPr:
            p = pos[etree.QName(c).localname]
            assert p >= last, 'pPr 元素顺序违规'
            last = p


@pytest.mark.pure
class TestSetFirstLineIndent:
    """set_first_line_indent 产物: 2 字符缩进 + 无互斥属性 + 属性顺序."""

    def test_attrs(self):
        doc = Document()
        p = doc.add_paragraph('正文')
        set_first_line_indent(p, chars=2)
        ind = p._p.find(qn('w:pPr') + '/' + qn('w:ind'))
        assert ind is not None
        assert ind.get(qn('w:firstLineChars')) == '200'
        # 互斥属性必须清除 (EG_HangingIndent choice)
        assert ind.get(qn('w:hangingChars')) is None
        assert ind.get(qn('w:hanging')) is None
        # 绝对数值清除, 避免与 firstLineChars 冲突
        assert ind.get(qn('w:firstLine')) is None
        assert ind.get(qn('w:left')) is None

    def test_attr_order(self):
        """ind 属性顺序符合 CT_Ind: firstLineChars 在 leftChars 前."""
        doc = Document()
        p = doc.add_paragraph('正文')
        set_first_line_indent(p, chars=2)
        ind = p._p.find(qn('w:pPr') + '/' + qn('w:ind'))
        attr_names = [etree.QName(k).localname for k in ind.attrib]
        assert attr_names.index('firstLineChars') < attr_names.index('leftChars'), \
            f'CT_Ind 属性顺序违规: {attr_names}'

    def test_rPr_exists_after_insert(self):
        """已有 rPr 的段落: ind 插入后 rPr 仍在其后 (无顺序违规)."""
        doc = Document()
        p = doc.add_paragraph('正文')
        rPr = p._p.find(qn('w:pPr') + '/' + qn('w:rPr'))
        if rPr is None:
            pPr = p._p.get_or_add_pPr()
            etree.SubElement(pPr, W + 'rPr')
        set_first_line_indent(p, chars=2)
        pPr = p._p.find(qn('w:pPr'))
        names = [etree.QName(c).localname for c in pPr]
        assert names.index('ind') < names.index('rPr'), names


@pytest.mark.pure
class TestSetLineSpacing:
    """set_line_spacing 产物: spacing 必须插在 jc 之前 (CT_PPr 顺序)."""

    def test_spacing_before_jc(self):
        """已有 jc 的段落加行距: spacing 不得 append 到 jc 后 (历史 bug)."""
        doc = Document()
        p = doc.add_paragraph('标题')
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER  # pPr 先有 jc
        set_line_spacing(p, 1.0)
        pPr = p._p.find(qn('w:pPr'))
        names = [etree.QName(c).localname for c in pPr]
        assert names.index('spacing') < names.index('jc'), \
            f'spacing 插到 jc 之后 → CT_PPr 顺序违规: {names}'


@pytest.mark.pure
class TestSeparatorType:
    """footnotes.xml 的 -1/0 分隔符必须带 w:type (否则 Word 拒开)."""

    def test_md2docx_build(self):
        fn_bytes = _build_footnotes_xml({1: '脚注文本'})
        root = etree.fromstring(fn_bytes)
        for fid, expected in [('-1', 'separator'), ('0', 'continuationSeparator')]:
            fn = next(f for f in root.iter(W + 'footnote')
                      if f.get(W + 'id') == fid)
            assert fn.get(W + 'type') == expected, \
                f'footnote id={fid} 缺少 w:type="{expected}" → Word 拒开'

    def test_conftest_inject(self, docx_with_footnotes):
        """conftest 注入的夹具文档也必须带 w:type."""
        with zipfile.ZipFile(docx_with_footnotes) as z:
            root = etree.fromstring(z.read('word/footnotes.xml'))
        for fid, expected in [('-1', 'separator'), ('0', 'continuationSeparator')]:
            fn = next(f for f in root.iter(W + 'footnote')
                      if f.get(W + 'id') == fid)
            assert fn.get(W + 'type') == expected

    def test_format_document_preserves(self, docx_with_footnotes, tmp_path):
        """format_document 输出: separator type 保持 + rel 单条."""
        out = str(tmp_path / 'out.docx')
        format_document(docx_with_footnotes, out, body_indent=2)

        with zipfile.ZipFile(out) as z:
            fn_root = etree.fromstring(z.read('word/footnotes.xml'))
            rels = z.read('word/_rels/document.xml.rels').decode()

        for fid, expected in [('-1', 'separator'), ('0', 'continuationSeparator')]:
            fn = next(f for f in fn_root.iter(W + 'footnote')
                      if f.get(W + 'id') == fid)
            assert fn.get(W + 'type') == expected
        # 重复 footnotes 关系 → Word 拒开
        n_rel = len(re.findall(r'<Relationship[^>]*footnotes', rels))
        assert n_rel == 1, f'footnotes 关系应为 1 条, 实际 {n_rel}'


@pytest.mark.integration
class TestFormatOutputValidity:
    """集成: format_document 输出整体静态合法性."""

    def test_full_static_scan(self, docx_with_footnotes, tmp_path):
        """输出文档: pPr 顺序 0 违规 + 正文缩进 2 字符."""
        out = str(tmp_path / 'out.docx')
        format_document(docx_with_footnotes, out, body_indent=2)

        with zipfile.ZipFile(out) as z:
            doc_root = etree.fromstring(z.read('word/document.xml'))

        # pPr 顺序违规 0 (ind 必须在 rPr 前)
        violations = 0
        for pPr in doc_root.iter(W + 'pPr'):
            names = [etree.QName(c).localname for c in pPr]
            pos = {name: i for i, name in enumerate(_CT_PPR)}
            last = -1
            for name in names:
                if name in ('bookmarkStart', 'bookmarkEnd'):  # 非 CT_PPr 但可插入
                    continue
                p = pos.get(name)
                if p is None:
                    continue
                if p < last:
                    violations += 1
                    break
                last = p
        assert violations == 0, 'pPr 元素顺序违规'

        # 正文段首行缩进 2 字符
        indented = 0
        for p in doc_root.iter(W + 'p'):
            ind = p.find(W + 'pPr/' + W + 'ind')
            if ind is not None and ind.get(W + 'firstLineChars') == '200':
                indented += 1
        assert indented >= 1, '正文段落缺少 firstLineChars=200'
