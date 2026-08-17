"""自定义格式规则测试: 规则加载/校验/字号表/自定义排版应用/四级标题.

覆盖: load_rules 合并与校验 / 中文字号名解析 / 固定磅值行距 /
自定义规则格式化端到端 / 四级标题检测与格式化 / roundtrip h4 /
脚注规则在 format_document 与 md2docx 两条路径生效 (回归: C11).
"""

import json
import re
import zipfile

import pytest
from docx import Document
from docx.oxml.ns import qn

from format_paper import detect_heading_level, format_document
from rules import (CN_SIZE_PT, DEFAULT_RULES, load_rules,
                   parse_line_spacing, parse_size)
from docx2md import docx_to_markdown
from md2docx import markdown_to_docx


def _write_rules(tmp_path, data) -> str:
    path = tmp_path / 'rules.json'
    path.write_text(json.dumps(data, ensure_ascii=False), encoding='utf-8')
    return str(path)


def _run_font(run):
    """Extract eastAsia font + size from a run."""
    rPr = run._r.find(qn('w:rPr'))
    ea = '?'
    if rPr is not None:
        rf = rPr.find(qn('w:rFonts'))
        if rf is not None:
            ea = rf.get(qn('w:eastAsia'), '?')
    return {'eastAsia': ea, 'size_pt': run.font.size.pt if run.font.size else 0,
            'bold': run.bold}


def _spacing(para):
    """Extract line spacing info: (line, rule)."""
    pPr = para._p.find(qn('w:pPr'))
    if pPr is None:
        return None
    sp = pPr.find(qn('w:spacing'))
    if sp is None:
        return None
    return (sp.get(qn('w:line')), sp.get(qn('w:lineRule')))


@pytest.mark.pure
class TestRuleParsing:
    """规则解析与校验."""

    def test_cn_size_table(self):
        """中文字号表关键值."""
        assert CN_SIZE_PT['二号'] == 22
        assert CN_SIZE_PT['四号'] == 14
        assert CN_SIZE_PT['小四'] == 12
        assert CN_SIZE_PT['五号'] == 10.5
        assert CN_SIZE_PT['小五'] == 9

    def test_parse_size_formats(self):
        """字号: 中文名/数字/磅值字符串."""
        assert parse_size('二号') == 22
        assert parse_size(22) == 22
        assert parse_size('22pt') == 22
        assert parse_size('22磅') == 22
        with pytest.raises(ValueError):
            parse_size('特大')

    def test_parse_line_spacing(self):
        """行距: 倍数/固定磅值."""
        assert parse_line_spacing(1.5) == 1.5
        assert parse_line_spacing('20磅') == {'mode': 'exact', 'value': 20.0}
        assert parse_line_spacing('20pt') == {'mode': 'exact', 'value': 20.0}
        assert parse_line_spacing('单倍') == 1.0
        with pytest.raises(ValueError):
            parse_line_spacing('abc')

    def test_load_rules_merge(self, tmp_path):
        """部分配置应与默认合并."""
        path = _write_rules(tmp_path, {'title': {'size': '二号', 'bold': True}})
        rules = load_rules(path)
        # 覆盖生效
        assert rules['title']['size'] == 22
        assert rules['title']['bold'] is True
        # 未写字段/元素用默认
        assert rules['title']['font'] == DEFAULT_RULES['title']['font']
        assert rules['body'] == DEFAULT_RULES['body']

    def test_load_rules_invalid_element(self, tmp_path):
        """未知元素应报错."""
        path = _write_rules(tmp_path, {'h5': {'size': 12}})
        with pytest.raises(ValueError, match='未知元素'):
            load_rules(path)

    def test_load_rules_comment_keys(self, tmp_path):
        """'_' 开头的键视为注释, 允许出现在配置中."""
        path = _write_rules(tmp_path, {'_说明': '这是注释', 'title': {'bold': True}})
        rules = load_rules(path)
        assert rules['title']['bold'] is True

    def test_load_rules_invalid_size(self, tmp_path):
        """非法字号应报错并指明元素."""
        path = _write_rules(tmp_path, {'title': {'size': '特大'}})
        with pytest.raises(ValueError, match='title'):
            load_rules(path)

    def test_load_rules_invalid_align(self, tmp_path):
        """非法对齐方式应报错."""
        path = _write_rules(tmp_path, {'body': {'align': 'diagonal'}})
        with pytest.raises(ValueError, match='对齐'):
            load_rules(path)

    def test_load_rules_missing_file(self):
        """文件不存在应报错."""
        with pytest.raises(FileNotFoundError):
            load_rules('不存在的.json')

    def test_default_rules_match_document_constants(self):
        """默认规则应与 skill 预设格式一致 (format_paper 常量)."""
        assert DEFAULT_RULES['title']['size'] == 14     # 四号
        assert DEFAULT_RULES['title']['font'] == '黑体'
        assert DEFAULT_RULES['h1']['size'] == 12        # 小四
        assert DEFAULT_RULES['h2']['font'] == '楷体'
        assert DEFAULT_RULES['body']['size'] == 10.5    # 五号
        assert DEFAULT_RULES['footnote']['size'] == 9   # 小五


@pytest.mark.integration
class TestCustomRuleFormatting:
    """自定义规则端到端排版."""

    @pytest.fixture
    def user_rules(self, tmp_path):
        """用户例子: 题目二号加粗 / 一级四号加粗 / 正文小四 20磅行距."""
        return load_rules(_write_rules(tmp_path, {
            'title': {'size': '二号', 'bold': True},
            'h1': {'font': '宋体', 'size': '四号', 'bold': True},
            'h2': {'font': '宋体', 'size': '小四', 'bold': True},
            'h3': {'font': '宋体', 'size': '小四', 'bold': False},
            'body': {'size': '小四', 'line_spacing': '20磅'},
        }))

    def test_user_example_formats(self, minimal_docx_path, tmp_path, user_rules):
        """用户格式要求应全部生效 (题目/一级/正文)."""
        out = str(tmp_path / 'out.docx')
        format_document(minimal_docx_path, out, rules=user_rules)

        doc = Document(out)
        paras = [p for p in doc.paragraphs if p.text.strip()]
        # 题目: 黑体二号(22)加粗
        info = _run_font(paras[0].runs[0])
        assert info['eastAsia'] == '黑体'
        assert info['size_pt'] == 22.0
        assert info['bold'] is True
        # 一级: 宋体四号(14)加粗
        for p in paras:
            if p.text.strip().startswith('一、'):
                info = _run_font(p.runs[0])
                assert info['eastAsia'] == '宋体'
                assert info['size_pt'] == 14.0
                assert info['bold'] is True
                break
        # 正文: 宋体小四(12) 20磅固定行距
        for p in paras:
            if '正文内容' in p.text:
                info = _run_font(p.runs[0])
                assert info['size_pt'] == 12.0
                assert info['bold'] is False
                assert _spacing(p) == ('400', 'exact')
                break

    def test_default_rules_unchanged(self, minimal_docx_path, tmp_path):
        """不传 rules 时默认行为不变 (题目黑体14不加粗)."""
        out = str(tmp_path / 'out.docx')
        format_document(minimal_docx_path, out)

        doc = Document(out)
        first = next(p for p in doc.paragraphs if p.text.strip())
        info = _run_font(first.runs[0])
        assert info['size_pt'] == 14.0
        assert info['bold'] is False

    def test_md2docx_custom_rules(self, tmp_path, user_rules):
        """md → docx 重建也应应用自定义规则."""
        md = '# 测试题目\n\n## 一、标题\n\n正文内容。\n'
        out = str(tmp_path / 'out.docx')
        markdown_to_docx(md, out, rules=user_rules)

        doc = Document(out)
        paras = [p for p in doc.paragraphs if p.text.strip()]
        info = _run_font(paras[0].runs[0])
        assert info['size_pt'] == 22.0 and info['bold'] is True
        body = next(p for p in paras if '正文' in p.text)
        assert _spacing(body) == ('400', 'exact')


def _footnotes_xml(path) -> str:
    """读取 docx 的 footnotes.xml 原文 (zipfile, 避免依赖解析库)."""
    with zipfile.ZipFile(path) as z:
        return z.read('word/footnotes.xml').decode('utf-8')


@pytest.mark.integration
class TestFootnoteRules:
    """脚注规则应在两条路径生效 (回归: C11 缺陷 — 规则未透传)."""

    @pytest.fixture
    def fn_rules(self, tmp_path):
        """脚注 1.5 倍行距 + 小五 (9pt)."""
        return load_rules(_write_rules(tmp_path, {
            'footnote': {'size': '小五', 'line_spacing': '1.5'},
        }))

    def test_format_document_applies_footnote_rules(
            self, tmp_path, fn_rules, docx_with_footnotes):
        """工作流 1 (format_document): 自定义脚注行距/字号生效."""
        out = str(tmp_path / 'out.docx')
        format_document(docx_with_footnotes, out, rules=fn_rules)

        xml = _footnotes_xml(out)
        assert 'w:line="360"' in xml, f'脚注应 1.5 倍行距 (line=360): {xml[:400]}'
        assert 'w:lineRule="auto"' in xml
        assert 'w:val="18"' in xml, f'脚注应 9pt (sz=18): {xml[:400]}'

    def test_md2docx_applies_footnote_rules(self, tmp_path, fn_rules):
        """md 中转路径 (markdown_to_docx): 自定义脚注行距/字号生效."""
        md = ('# 测试题目\n\n正文引用脚注[^1]。\n\n'
              '[^1]: 参见作者：《书名》，载《期刊》2020年第1期，第5页。\n')
        out = str(tmp_path / 'out.docx')
        markdown_to_docx(md, out, rules=fn_rules)

        xml = _footnotes_xml(out)
        assert 'w:line="360"' in xml, f'脚注应 1.5 倍行距 (line=360): {xml[:400]}'
        assert 'w:val="18"' in xml, f'脚注应 9pt (sz=18): {xml[:400]}'
        assert 'w:eastAsia="宋体"' in xml, f'脚注应宋体: {xml[:400]}'


@pytest.mark.pure
class TestLevel4Detection:
    """四级标题 (（1）) 检测."""

    def test_level4_recognized(self):
        """（1）短文本 → 4."""
        assert detect_heading_level('（1）直接歧视') == 4

    def test_level4_long_not_heading(self):
        """超长 （1）段 → 正文 (避免误判)."""
        assert detect_heading_level('（1）合同解除权与违约责任制度的衔接适用分析') == 0

    def test_level3_unchanged(self):
        """三级标题检测不受影响."""
        assert detect_heading_level('1. 研究方法') == 3

    def test_level2_unchanged(self):
        """二级标题 (中文数字括号) 不受影响."""
        assert detect_heading_level('（一）研究背景') == 2


@pytest.mark.integration
class TestLevel4Roundtrip:
    """四级标题 docx → md → docx 往返."""

    def test_h4_roundtrip(self, tmp_path):
        """（1）标题: docx → md (##### 前缀) → docx (h4 规则)."""
        doc = Document()
        doc.add_paragraph('四级标题测试')
        doc.add_paragraph('一、总述')
        doc.add_paragraph('（1）直接歧视')
        doc.add_paragraph('正文内容。')
        inp = str(tmp_path / 'in.docx')
        doc.save(inp)

        md = docx_to_markdown(inp)
        assert '\n##### （1）直接歧视' in md, f'md 中应为 ##### 前缀:\n{md}'

        out = str(tmp_path / 'out.docx')
        stats = markdown_to_docx(md, out)
        assert stats['headings_l4'] == 1

        result = Document(out)
        h4 = next(p for p in result.paragraphs
                  if p.text.strip().startswith('（1）'))
        info = _run_font(h4.runs[0])
        # 默认 h4 规则: 宋体五号加粗缩进2
        assert info['eastAsia'] == '宋体'
        assert info['size_pt'] == 10.5
        assert info['bold'] is True


class TestPunctuationSkipHeadings:
    """标点规范化跳过题目与各级标题 (回归: 用户反馈标题标点/空格被改).

    修复前: 标点规范化遍历所有段落, 标题 "一、 绪论" 的空格会被空格清理
    规则删除, run 拆分的 "1." 会被终端句号规则改写成 "1。"。
    """

    def test_format_document_skips_headings(self, tmp_path):
        """原地路径: 题目/标题标点不动, 摘要与正文仍规范化."""
        doc = Document()
        doc.add_paragraph('测试论文')
        doc.add_paragraph('【摘要】本文研究了,相关问题。')
        doc.add_paragraph('一、 绪论')
        doc.add_paragraph('1. 理论意义')
        doc.add_paragraph('这是正文内容,包含英文逗号。')
        path = str(tmp_path / 'headings.docx')
        doc.save(path)
        out = str(tmp_path / 'out.docx')
        format_document(path, out, body_indent=2)

        result = Document(out)
        texts = [p.text for p in result.paragraphs]
        assert texts[0] == '测试论文'                          # 题目未动
        assert texts[1] == '【摘要】本文研究了，相关问题。'      # 摘要仍规范化
        assert texts[2] == '一、 绪论'                         # 标题空格未被清理
        assert texts[3] == '1. 理论意义'                       # 三级标题编号标点未动
        assert texts[4] == '这是正文内容，包含英文逗号。'        # 正文逗号已规范化

    def test_markdown_to_docx_skips_headings(self, tmp_path):
        """中转路径: 重建后题目/标题标点不动, 正文仍规范化."""
        md = '\n'.join([
            '# 测试论文',
            '## 一、 绪论',
            '#### 1. 理论意义',
            '这是正文内容,包含英文逗号。',
            '',
        ])
        out = str(tmp_path / 'out.docx')
        markdown_to_docx(md, out)

        result = Document(out)
        texts = [p.text for p in result.paragraphs]
        assert texts[0] == '测试论文'                          # 题目未动
        assert texts[1] == '一、 绪论'                         # 标题空格未被清理
        assert texts[2] == '1. 理论意义'                       # 三级标题编号标点未动
        assert texts[3] == '这是正文内容，包含英文逗号。'        # 正文逗号已规范化
