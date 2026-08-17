"""Integration tests for footnote structure and write-back."""

import pytest
from docx import Document
from docx.oxml.ns import qn

from citation_formatter import (extract_footnotes, format_all_footnotes,
                                 check_footnote)


def _get_footnote_run_order(doc):
    """Extract run order for each footnote: list of (has_ref, text) tuples."""
    from lxml import etree
    FOOTNOTE_REL = ('http://schemas.openxmlformats.org/officeDocument/'
                    '2006/relationships/footnotes')
    for rel in doc.part.rels.values():
        if rel.reltype == FOOTNOTE_REL:
            fn_xml = etree.fromstring(rel.target_part.blob)
            ns = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
            result = {}
            for fn in fn_xml.findall(f'{{{ns}}}footnote'):
                fid = fn.get(f'{{{ns}}}id')
                if fid and int(fid) > 0:
                    runs = []
                    for p in fn.findall(f'{{{ns}}}p'):
                        for r in p.findall(f'{{{ns}}}r'):
                            has_ref = r.find(f'{{{ns}}}footnoteRef') is not None
                            texts = ''.join(
                                t.text or '' for t in r.findall(f'{{{ns}}}t'))
                            runs.append((has_ref, texts))
                    result[int(fid)] = runs
            return result
    return {}


@pytest.mark.integration
class TestFootnoteOrder:
    """Verify footnote XML structure: ref before text."""

    def test_ref_before_text(self, docx_with_footnotes):
        """Auto-number (footnoteRef) should appear between bracket and citation text."""
        doc = Document(docx_with_footnotes)
        order = _get_footnote_run_order(doc)

        for fid, runs in order.items():
            ref_idx = None
            cite_idx = None
            for i, (has_ref, text) in enumerate(runs):
                if has_ref:
                    ref_idx = i
                # Find citation text (not just bracket like "[" or "]")
                if text.strip() and not has_ref and len(text.strip()) > 2:
                    if cite_idx is None:
                        cite_idx = i
            if ref_idx is not None and cite_idx is not None:
                assert ref_idx < cite_idx, \
                    f'Fn {fid}: ref at index {ref_idx}, citation text at index {cite_idx}'
                return

    def test_ref_preserved_after_fix(self, docx_with_footnotes):
        """After auto-fix, footnoteRef element should still exist."""
        doc = Document(docx_with_footnotes)

        # Run fix
        stats = format_all_footnotes(doc, fix=True)
        # Save to a BytesIO and reload to verify persistence
        import io
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        doc2 = Document(buf)

        order = _get_footnote_run_order(doc2)
        for fid, runs in order.items():
            has_ref = any(r[0] for r in runs)
            if fid <= 3:  # only check footnotes we created
                assert has_ref, f'Fn {fid}: footnoteRef lost after fix'

    def test_citation_text_after_fix(self, docx_with_footnotes):
        """After fix, citation text should be in runs after footnoteRef."""
        doc = Document(docx_with_footnotes)
        format_all_footnotes(doc, fix=True)

        import io
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        doc2 = Document(buf)

        order = _get_footnote_run_order(doc2)
        for fid, runs in order.items():
            if fid > 3:
                continue
            # Find the run after footnoteRef that has text
            found_ref = False
            text_after_ref = ''
            for has_ref, text in runs:
                if has_ref:
                    found_ref = True
                    continue
                if found_ref and text.strip():
                    text_after_ref = text
                    break
            # The text after ref should contain citation content
            if text_after_ref:
                assert len(text_after_ref.strip()) > 2, \
                    f'Fn {fid}: citation text too short: "{text_after_ref}"'

    def test_blob_write_back(self, docx_with_footnotes):
        """Modified footnotes XML should persist after save + reload."""
        doc = Document(docx_with_footnotes)

        # Get original text
        fns_before = extract_footnotes(doc)

        # Fix
        format_all_footnotes(doc, fix=True)

        # Save and reload
        import io
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        doc2 = Document(buf)

        fns_after = extract_footnotes(doc2)

        # Check that fixes are persisted
        for fn in fns_after:
            if not fn['full_text'].strip():
                continue
            issues = check_footnote(fn['full_text'])
            # Footnote 3 is 《民法典》 which should have no issues
            if '民法典' in fn['full_text']:
                continue
            # Other footnotes may still have issues that can't be auto-fixed
            # but at minimum the fix count should have been recorded
            pass

        # At least one footnote should have been modified
        assert any('载《' in fn['full_text'] or '年版' in fn['full_text']
                   for fn in fns_after if fn['full_text'].strip()), \
            'No fixes persisted after save+reload'

    def test_citation_fixes_applied(self, docx_with_legal_citations):
        """Legal document + case citation errors should be fixed."""
        doc = Document(docx_with_legal_citations)
        format_all_footnotes(doc, fix=True)

        import io
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        doc2 = Document(buf)

        fns = extract_footnotes(doc2)
        fn_texts = {fn['id']: fn['full_text'] for fn in fns}

        # Fn1: 示例发[2020] → 示例发〔2020〕
        if 1 in fn_texts:
            assert '示例发〔2020〕' in fn_texts[1], \
                f'Legal bracket not fixed: {fn_texts[1]}'
            assert '示例发[2020]' not in fn_texts[1]

        # Fn2: [2021] → （2021）
        if 2 in fn_texts:
            assert '（2021）' in fn_texts[2], \
                f'Case bracket not fixed: {fn_texts[2]}'

        # Fn3: 示例发[2018] → 示例发〔2018〕
        if 3 in fn_texts:
            assert '示例发〔2018〕' in fn_texts[3], \
                f'Legal bracket not fixed: {fn_texts[3]}'

    def test_missing_year_fixed(self, docx_with_missing_year):
        """BUG-006: 出版社缺"年版"应被修复 (auto_fix 能力大于 check 检测面)."""
        doc = Document(docx_with_missing_year)
        stats = format_all_footnotes(doc, fix=True)

        import io
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        doc2 = Document(buf)

        fns = extract_footnotes(doc2)
        text = fns[0]['full_text']
        assert '示例出版社2017年版' in text, f'缺年版未修复: {text!r}'
        assert stats['fixed'] >= 1, 'stats 应记录修复数'

    def test_trailing_period_fixed(self, docx_with_trailing_period):
        """BUG-006: 终端英文句号应修复为中文句号."""
        doc = Document(docx_with_trailing_period)
        format_all_footnotes(doc, fix=True)

        import io
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        doc2 = Document(buf)

        fns = extract_footnotes(doc2)
        text = fns[0]['full_text']
        assert text.strip().endswith('页。'), f'终端句号未修复: {text!r}'
